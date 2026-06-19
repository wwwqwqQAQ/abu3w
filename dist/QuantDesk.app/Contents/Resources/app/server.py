#!/usr/bin/env python3
"""QuantDesk — 量化交易工作站后端"""
import subprocess, json, math, os, sys, io, csv, time, asyncio, hashlib, re, threading
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from fastapi import FastAPI, Query, WebSocket, WebSocketDisconnect
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import warnings; warnings.filterwarnings('ignore')

sys.path.insert(0, os.path.dirname(__file__))
app = FastAPI(title="QuantDesk")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# ═══════════ 数据 ═══════════
KLINE_CACHE = {}
KLINE_CACHE_TTL = 60 * 60 * 24
KLINE_CACHE_DIR = os.path.expanduser("~/Library/Application Support/QuantDesk/cache/kline")
KLINE_CONNECT_TIMEOUT = "1"
KLINE_MAX_TIME = "3"
KLINE_PROCESS_TIMEOUT = 4


def cache_key(code, days):
    return f"{code}_{int(days)}"


def cache_path(code, days):
    return os.path.join(KLINE_CACHE_DIR, f"{cache_key(code, days)}.json")


def read_kline_cache(code, days, allow_stale=False):
    now = time.time()
    for cached_days in sorted({int(days), 500, 300, 200, 100}, reverse=True):
        if cached_days < int(days):
            continue
        key = cache_key(code, cached_days)
        item = KLINE_CACHE.get(key)
        if item and (allow_stale or now - item["time"] <= KLINE_CACHE_TTL):
            return item["data"][-int(days):]

        path = cache_path(code, cached_days)
        try:
            if not os.path.exists(path):
                continue
            if not allow_stale and now - os.path.getmtime(path) > KLINE_CACHE_TTL:
                continue
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            KLINE_CACHE[key] = {"time": os.path.getmtime(path), "data": data}
            return data[-int(days):]
        except Exception:
            continue
    return None


def write_kline_cache(code, days, data):
    if not data:
        return
    key = cache_key(code, days)
    KLINE_CACHE[key] = {"time": time.time(), "data": data}
    try:
        os.makedirs(KLINE_CACHE_DIR, exist_ok=True)
        with open(cache_path(code, days), "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, separators=(",", ":"))
    except Exception:
        pass


def load_anthropic_api_key():
    key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if key:
        return key

    support_dir = os.path.expanduser("~/Library/Application Support/QuantDesk")
    json_paths = [
        os.path.join(support_dir, "config.json"),
        os.path.expanduser("~/.quantdesk.json"),
    ]
    env_paths = [
        os.path.join(support_dir, ".env"),
        os.path.join(support_dir, "anthropic_api_key.txt"),
        os.path.expanduser("~/.quantdesk.env"),
    ]

    for path in json_paths:
        try:
            if not os.path.exists(path):
                continue
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            key = str(data.get("ANTHROPIC_API_KEY") or data.get("anthropic_api_key") or "").strip()
            if key:
                return key
        except Exception:
            continue

    for path in env_paths:
        try:
            if not os.path.exists(path):
                continue
            with open(path, "r", encoding="utf-8") as f:
                text = f.read().strip()
            if path.endswith(".txt") and "=" not in text:
                return text.splitlines()[0].strip()
            for line in text.splitlines():
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                name, value = line.split("=", 1)
                if name.strip() == "ANTHROPIC_API_KEY":
                    return value.strip().strip("\"'")
        except Exception:
            continue

    return ""


def fetch_kline(code, days=500):
    cached = read_kline_cache(code, days)
    if cached is not None:
        return cached

    m = 'sh' if code.startswith(('6','9')) else 'sz'
    url = f'http://money.finance.sina.com.cn/quotes_service/api/json_v2.php/CN_MarketData.getKLineData?symbol={m}{code}&scale=240&ma=no&datalen={days}'
    try:
        r = subprocess.run(
            ['curl','-sS','--connect-timeout',KLINE_CONNECT_TIMEOUT,'--max-time',KLINE_MAX_TIME,'--noproxy','*',url],
            capture_output=True,
            text=True,
            timeout=KLINE_PROCESS_TIMEOUT
        )
        if r.returncode != 0 or not r.stdout or r.stdout.startswith('null'):
            return read_kline_cache(code, days, allow_stale=True) or []
        data = [(k['day'],float(k['open']),float(k['close']),float(k['high']),float(k['low']),int(float(k['volume']))) for k in json.loads(r.stdout)]
        write_kline_cache(code, days, data)
        return data
    except Exception:
        return read_kline_cache(code, days, allow_stale=True) or []

def fetch_realtime(codes):
    if not codes: return {}
    ids = ','.join(codes)
    try:
        r = subprocess.run(
            ['curl','-sS','--connect-timeout','1','--max-time','3','--noproxy','*','-H','Referer: https://finance.sina.com.cn','-H','User-Agent: Mozilla/5.0',f'http://hq.sinajs.cn/list={ids}'],
            capture_output=True,
            timeout=4
        )
        if r.returncode != 0:
            return {}
        text = r.stdout.decode('gbk', errors='replace')
    except Exception:
        return {}
    results = {}
    for line in text.strip().split('\n'):
        if not line or '=' not in line: continue
        c = line.split('=')[0].split('_')[-1]
        raw = line.split('"')[1] if '"' in line else ''
        if raw:
            p = raw.split(',')
            if len(p)>=6:
                try: results[c] = {'name':p[0],'open':float(p[1]),'yest_close':float(p[2]),'price':float(p[3]),'high':float(p[4]),'low':float(p[5]),'volume':0}
                except: pass
    return results

# ═══════════ 指标 ═══════════
def calc_rsi(closes, n=14):
    if len(closes)<n+1: return 50
    g=[max(closes[i]-closes[i-1],0) for i in range(1,len(closes))]
    l=[max(closes[i-1]-closes[i],0) for i in range(1,len(closes))]
    ag=sum(g[-n:])/n; al=sum(l[-n:])/n
    if al==0: return 100
    return 100-100/(1+ag/(al+0.0001))

def calc_ma(prices,n):
    if len(prices)<n: return sum(prices)/len(prices)
    return sum(prices[-n:])/n

# ═══════════ 增强回测 ═══════════
def run_backtest_detailed(code, params):
    """详细回测 — 返回每笔交易明细 + 风险指标"""
    kl = fetch_kline(code, 500)
    if not kl: return None

    o=[k[1] for k in kl]; c=[k[2] for k in kl]; h=[k[3] for k in kl]; l=[k[4] for k in kl]
    trades=[]; holding=False; bp=0; bd=''; cur=100000; max_price=0; max_dd=0; peak=100000

    sid = params.get("strategy","rsi")
    # 构建 RSI
    r=[50]*14
    for i in range(1,len(c)):
        g=[max(c[j]-c[j-1],0) for j in range(1,i+1)]
        ls=[max(c[j-1]-c[j],0) for j in range(1,i+1)]
        ag=sum(g[-14:])/14; al=sum(ls[-14:])/14
        r.append(100-100/(1+ag/(al+0.0001)))

    # 计算均线
    def ma_at(prices,n,i):
        if i<n: return sum(prices[:i+1])/(i+1)
        return sum(prices[i-n+1:i+1])/n

    for i in range(60,len(kl)-1):
        # 更新峰值
        peak = max(peak, cur)
        dd = (peak-cur)/peak if peak>0 else 0
        max_dd = max(max_dd, dd)

        if not holding:
            should_buy = False

            if sid=="rsi":
                os=params.get("oversold",25); should_buy=r[i]<os
            elif sid=="ma_cross":
                f=params.get("fast",5); s=params.get("slow",20)
                mf=ma_at(c,f,i); mf_prev=ma_at(c,f,i-1)
                ms=ma_at(c,s,i); ms_prev=ma_at(c,s,i-1)
                should_buy = mf_prev<=ms_prev and mf>ms
            elif sid=="turtle":
                e=params.get("entry",20); should_buy=h[i]>=max(h[max(0,i-e):i])
            elif sid=="momentum":
                mp=params.get("ma_period",60); mt=params.get("mom_threshold",0.10)
                mom=(c[i]-c[i-20])/c[i-20] if i>=20 else 0
                should_buy = c[i]>ma_at(c,mp,i) and mom>mt and r[i]>60
            elif sid=="rsi_trailing":
                os=params.get("oversold",25); should_buy=r[i]<os
            elif sid=="custom":
                # 自定义表达式求值
                should_buy = eval_condition(params.get("buy_condition","False"), c,o,h,l,r,i)

            if should_buy:
                bp=o[i+1]; bd=kl[i+1][0]; max_price=bp; holding=True

        elif holding:
            should_sell = False
            max_price = max(max_price, h[i])

            if sid=="rsi":
                ob=params.get("overbought",80); should_sell=r[i]>ob
            elif sid=="ma_cross":
                f=params.get("fast",5); s=params.get("slow",20)
                should_sell = ma_at(c,f,i)<=ma_at(c,s,i)
            elif sid=="turtle":
                x=params.get("exit_",10); should_sell=l[i]<=min(l[max(0,i-x):i])
            elif sid=="momentum":
                should_sell = c[i]<ma_at(c,20,i) or r[i]<40
            elif sid=="rsi_trailing":
                tp=params.get("trail_pct",0.08); re=params.get("rsi_exit",75)
                should_sell = c[i]<=max_price*(1-tp) or r[i]>re
            elif sid=="custom":
                should_sell = eval_condition(params.get("sell_condition","False"), c,o,h,l,r,i)

            if should_sell:
                sp = o[i+1]; sd = kl[i+1][0]
                sh = int(cur/bp/100)*100
                if sh>=100:
                    days = i - next((j for j in range(len(kl)) if kl[j][0]==bd), i)
                    pnl=(sp-bp)*sh; cur+=pnl
                    trades.append({
                        'buy_date':bd,'buy_price':round(bp,2),'sell_date':sd,'sell_price':round(sp,2),
                        'pnl':round(pnl,2),'pct':round((sp-bp)/bp*100,2),
                        'holding_days':days,'peak_dd':round((max_price-bp)/bp*100,2)
                    })
                holding=False

    if holding:
        sp=c[-1]; sd=kl[-1][0]
        sh=int(cur/bp/100)*100
        if sh>=100:
            days=len(kl)-1-next((j for j in range(len(kl)) if kl[j][0]==bd),0)
            pnl=(sp-bp)*sh; cur+=pnl
            trades.append({
                'buy_date':bd,'buy_price':round(bp,2),'sell_date':sd,'sell_price':round(sp,2),
                'pnl':round(pnl,2),'pct':round((sp-bp)/bp*100,2),
                'holding_days':days,'peak_dd':round((max_price-bp)/bp*100,2)
            })

    total_pnl=cur-100000
    wins=sum(1 for t in trades if t['pnl']>0)
    total=len(trades)

    # Sharpe ratio (年化)
    if total>=3:
        returns=[t['pct']/100 for t in trades]
        avg_ret=sum(returns)/len(returns)
        std_ret=(sum((r-avg_ret)**2 for r in returns)/len(returns))**0.5
        sharpe = round(avg_ret/std_ret*(total**0.5),2) if std_ret>0 else 0
    else:
        sharpe=0

    # 盈亏比
    avg_win=sum(t['pct'] for t in trades if t['pnl']>0)/max(1,sum(1 for t in trades if t['pnl']>0))
    avg_loss=sum(t['pct'] for t in trades if t['pnl']<0)/max(1,sum(1 for t in trades if t['pnl']<0))

    return {
        'pnl':total_pnl,'trades':total,'wins':wins,
        'win_rate':round(wins/total*100,1) if total else 0,
        'avg_pct':round(sum(t['pct'] for t in trades)/total,2) if total else 0,
        'sharpe':sharpe,'max_drawdown':round(max_dd*100,2),
        'avg_win':round(avg_win,2),'avg_loss':round(avg_loss,2),
        'profit_factor':round(abs(avg_win/avg_loss),1) if avg_loss!=0 else 99,
        'trade_list':trades,
    }

def eval_condition(expr, c, o, h, l, rsi_vals, i):
    """安全求值用户自定义条件表达式"""
    if not expr or expr.strip()=='False': return False
    try:
        local_vars = {
            'close': c[i], 'open': o[i], 'high': h[i], 'low': l[i],
            'rsi_14': rsi_vals[i],
            'prev_close': c[i-1] if i>0 else c[i],
            'ma_20': sum(c[max(0,i-19):i+1])/min(20,i+1),
            'ma_60': sum(c[max(0,i-59):i+1])/min(60,i+1),
            'ma_120': sum(c[max(0,i-119):i+1])/min(120,i+1),
            'change_5d': (c[i]-c[i-5])/c[i-5]*100 if i>=5 else 0,
            'change_20d': (c[i]-c[i-20])/c[i-20]*100 if i>=20 else 0,
            'vol_20': (sum((c[j]-sum(c[max(0,i-19):i+1])/min(20,i+1))**2 for j in range(max(0,i-19),i+1))/min(20,i+1))**0.5/c[i]*100 if i>=19 else 1,
            'math': math,
        }
        return bool(eval(expr, {"__builtins__": {}}, local_vars))
    except Exception:
        return False

# ═══════════ 股票池 ═══════════
STOCK_POOL = [
    {"code":"600519","name":"茅台","market":"sh"},{"code":"600036","name":"招行","market":"sh"},
    {"code":"601318","name":"中国平安","market":"sh"},{"code":"600276","name":"恒瑞医药","market":"sh"},
    {"code":"600900","name":"长江电力","market":"sh"},{"code":"601166","name":"兴业银行","market":"sh"},
    {"code":"600030","name":"中信证券","market":"sh"},{"code":"600887","name":"伊利股份","market":"sh"},
    {"code":"601398","name":"工商银行","market":"sh"},{"code":"601899","name":"紫金矿业","market":"sh"},
    {"code":"603259","name":"药明康德","market":"sh"},{"code":"601012","name":"隆基绿能","market":"sh"},
    {"code":"601888","name":"中国中免","market":"sh"},{"code":"600809","name":"山西汾酒","market":"sh"},
    {"code":"601857","name":"中国石油","market":"sh"},{"code":"600028","name":"中国石化","market":"sh"},
    {"code":"601288","name":"农业银行","market":"sh"},{"code":"600690","name":"海尔智家","market":"sh"},
    {"code":"000001","name":"平安银行","market":"sz"},{"code":"000858","name":"五粮液","market":"sz"},
    {"code":"002415","name":"海康威视","market":"sz"},{"code":"300750","name":"宁德时代","market":"sz"},
    {"code":"000333","name":"美的集团","market":"sz"},{"code":"002594","name":"比亚迪","market":"sz"},
    {"code":"000568","name":"泸州老窖","market":"sz"},{"code":"300059","name":"东方财富","market":"sz"},
    {"code":"002475","name":"立讯精密","market":"sz"},{"code":"000725","name":"京东方","market":"sz"},
    {"code":"002714","name":"牧原股份","market":"sz"},{"code":"300015","name":"爱尔眼科","market":"sz"},
    {"code":"000063","name":"中兴通讯","market":"sz"},{"code":"002352","name":"顺丰控股","market":"sz"},
    {"code":"000002","name":"万科A","market":"sz"},{"code":"002304","name":"洋河股份","market":"sz"},
    {"code":"601668","name":"中国建筑","market":"sh"},{"code":"601988","name":"中国银行","market":"sh"},
    {"code":"601328","name":"交通银行","market":"sh"},{"code":"600000","name":"浦发银行","market":"sh"},
    {"code":"600048","name":"保利发展","market":"sh"},{"code":"600050","name":"中国联通","market":"sh"},
    {"code":"600104","name":"上汽集团","market":"sh"},{"code":"600111","name":"北方稀土","market":"sh"},
    {"code":"600150","name":"中国船舶","market":"sh"},{"code":"600196","name":"复星医药","market":"sh"},
    {"code":"600309","name":"万华化学","market":"sh"},{"code":"600406","name":"国电南瑞","market":"sh"},
    {"code":"600436","name":"片仔癀","market":"sh"},{"code":"600438","name":"通威股份","market":"sh"},
    {"code":"600489","name":"中金黄金","market":"sh"},{"code":"600547","name":"山东黄金","market":"sh"},
    {"code":"600570","name":"恒生电子","market":"sh"},{"code":"600585","name":"海螺水泥","market":"sh"},
    {"code":"600660","name":"福耀玻璃","market":"sh"},{"code":"600674","name":"川投能源","market":"sh"},
    {"code":"600745","name":"闻泰科技","market":"sh"},{"code":"600760","name":"中航沈飞","market":"sh"},
    {"code":"600795","name":"国电电力","market":"sh"},{"code":"600837","name":"海通证券","market":"sh"},
    {"code":"600845","name":"宝信软件","market":"sh"},{"code":"600893","name":"航发动力","market":"sh"},
    {"code":"600905","name":"三峡能源","market":"sh"},{"code":"600941","name":"中国移动","market":"sh"},
    {"code":"600958","name":"东方证券","market":"sh"},{"code":"600999","name":"招商证券","market":"sh"},
    {"code":"601006","name":"大秦铁路","market":"sh"},{"code":"601009","name":"南京银行","market":"sh"},
    {"code":"601088","name":"中国神华","market":"sh"},{"code":"601100","name":"恒立液压","market":"sh"},
    {"code":"601111","name":"中国国航","market":"sh"},{"code":"601138","name":"工业富联","market":"sh"},
    {"code":"601155","name":"新城控股","market":"sh"},{"code":"601186","name":"中国铁建","market":"sh"},
    {"code":"601211","name":"国泰君安","market":"sh"},{"code":"601225","name":"陕西煤业","market":"sh"},
    {"code":"601229","name":"上海银行","market":"sh"},{"code":"601997","name":"贵阳银行","market":"sh"},
    {"code":"601319","name":"中国人保","market":"sh"},{"code":"601336","name":"新华保险","market":"sh"},
    {"code":"601360","name":"三六零","market":"sh"},{"code":"601390","name":"中国中铁","market":"sh"},
    {"code":"601601","name":"中国太保","market":"sh"},{"code":"601628","name":"中国人寿","market":"sh"},
    {"code":"601633","name":"长城汽车","market":"sh"},{"code":"601688","name":"华泰证券","market":"sh"},
    {"code":"601728","name":"中国电信","market":"sh"},{"code":"601766","name":"中国中车","market":"sh"},
    {"code":"601800","name":"中国交建","market":"sh"},{"code":"601816","name":"京沪高铁","market":"sh"},
    {"code":"601818","name":"光大银行","market":"sh"},{"code":"601838","name":"成都银行","market":"sh"},
    {"code":"601868","name":"中国能建","market":"sh"},{"code":"601872","name":"招商轮船","market":"sh"},
    {"code":"601919","name":"中远海控","market":"sh"},{"code":"601939","name":"建设银行","market":"sh"},
    {"code":"601985","name":"中国核电","market":"sh"},{"code":"601989","name":"中国重工","market":"sh"},
    {"code":"603288","name":"海天味业","market":"sh"},{"code":"603501","name":"韦尔股份","market":"sh"},
    {"code":"603799","name":"华友钴业","market":"sh"},{"code":"603986","name":"兆易创新","market":"sh"},
    {"code":"688008","name":"澜起科技","market":"sh"},{"code":"688012","name":"中微公司","market":"sh"},
    {"code":"688036","name":"传音控股","market":"sh"},{"code":"688111","name":"金山办公","market":"sh"},
    {"code":"688126","name":"沪硅产业","market":"sh"},{"code":"688169","name":"石头科技","market":"sh"},
    {"code":"688187","name":"时代电气","market":"sh"},{"code":"688223","name":"晶科能源","market":"sh"},
    {"code":"000066","name":"中国长城","market":"sz"},{"code":"000069","name":"华侨城A","market":"sz"},
    {"code":"000100","name":"TCL科技","market":"sz"},{"code":"000157","name":"中联重科","market":"sz"},
    {"code":"000338","name":"潍柴动力","market":"sz"},{"code":"000425","name":"徐工机械","market":"sz"},
    {"code":"000538","name":"云南白药","market":"sz"},{"code":"000625","name":"长安汽车","market":"sz"},
    {"code":"000651","name":"格力电器","market":"sz"},{"code":"000661","name":"长春高新","market":"sz"},
    {"code":"000768","name":"中航西飞","market":"sz"},{"code":"000776","name":"广发证券","market":"sz"},
    {"code":"000786","name":"北新建材","market":"sz"},{"code":"000792","name":"盐湖股份","market":"sz"},
    {"code":"000800","name":"一汽解放","market":"sz"},{"code":"000876","name":"新希望","market":"sz"},
    {"code":"000895","name":"双汇发展","market":"sz"},{"code":"000963","name":"华东医药","market":"sz"},
    {"code":"001979","name":"招商蛇口","market":"sz"},{"code":"002001","name":"新和成","market":"sz"},
    {"code":"002007","name":"华兰生物","market":"sz"},{"code":"002027","name":"分众传媒","market":"sz"},
    {"code":"002049","name":"紫光国微","market":"sz"},{"code":"002050","name":"三花智控","market":"sz"},
    {"code":"002129","name":"TCL中环","market":"sz"},{"code":"002179","name":"中航光电","market":"sz"},
    {"code":"002230","name":"科大讯飞","market":"sz"},{"code":"002236","name":"大华股份","market":"sz"},
    {"code":"002241","name":"歌尔股份","market":"sz"},{"code":"002271","name":"东方雨虹","market":"sz"},
    {"code":"002311","name":"海大集团","market":"sz"},{"code":"002371","name":"北方华创","market":"sz"},
    {"code":"002410","name":"广联达","market":"sz"},{"code":"002460","name":"赣锋锂业","market":"sz"},
    {"code":"002466","name":"天齐锂业","market":"sz"},{"code":"002493","name":"荣盛石化","market":"sz"},
    {"code":"002555","name":"三七互娱","market":"sz"},{"code":"002601","name":"龙佰集团","market":"sz"},
    {"code":"002709","name":"天赐材料","market":"sz"},{"code":"002812","name":"恩捷股份","market":"sz"},
    {"code":"002916","name":"深南电路","market":"sz"},{"code":"002920","name":"德赛西威","market":"sz"},
    {"code":"300014","name":"亿纬锂能","market":"sz"},{"code":"300122","name":"智飞生物","market":"sz"},
    {"code":"300124","name":"汇川技术","market":"sz"},{"code":"300274","name":"阳光电源","market":"sz"},
    {"code":"300308","name":"中际旭创","market":"sz"},{"code":"300316","name":"晶盛机电","market":"sz"},
    {"code":"300347","name":"泰格医药","market":"sz"},{"code":"300408","name":"三环集团","market":"sz"},
    {"code":"300413","name":"芒果超媒","market":"sz"},{"code":"300433","name":"蓝思科技","market":"sz"},
    {"code":"300450","name":"先导智能","market":"sz"},{"code":"300454","name":"深信服","market":"sz"},
    {"code":"300496","name":"中科创达","market":"sz"},{"code":"300498","name":"温氏股份","market":"sz"},
    {"code":"300628","name":"亿联网络","market":"sz"},{"code":"300760","name":"迈瑞医疗","market":"sz"},
    {"code":"300782","name":"卓胜微","market":"sz"},{"code":"300896","name":"爱美客","market":"sz"},
    {"code":"300919","name":"中伟股份","market":"sz"},{"code":"300957","name":"贝泰妮","market":"sz"},
]

CACHE_WARMUP_STARTED = False


def warm_kline_cache():
    global CACHE_WARMUP_STARTED
    if CACHE_WARMUP_STARTED:
        return
    CACHE_WARMUP_STARTED = True
    codes = [s["code"] for s in STOCK_POOL]

    def warm_one(code):
        if read_kline_cache(code, 500, allow_stale=True):
            return
        fetch_kline(code, 500)

    with ThreadPoolExecutor(max_workers=16) as pool:
        for _ in as_completed([pool.submit(warm_one, code) for code in codes]):
            pass


@app.on_event("startup")
def start_cache_warmup():
    threading.Thread(target=warm_kline_cache, daemon=True).start()


PRESET_STRATEGIES = {
    "rsi":{"name":"RSI反转","desc":"RSI超卖买入，超买卖出","params":{"oversold":25,"overbought":80},
        "param_meta":{"oversold":{"type":"range","min":15,"max":35,"desc":"超卖阈值"},"overbought":{"type":"range","min":65,"max":85,"desc":"超买阈值"}}},
    "ma_cross":{"name":"均线金叉","desc":"快线上穿慢线买，下穿卖","params":{"fast":5,"slow":20},
        "param_meta":{"fast":{"type":"range","min":3,"max":15,"desc":"快线周期"},"slow":{"type":"range","min":15,"max":90,"desc":"慢线周期"}}},
    "turtle":{"name":"海龟突破","desc":"突破N日高买，跌破M日低卖","params":{"entry":20,"exit_":10},
        "param_meta":{"entry":{"type":"range","min":10,"max":60,"desc":"突破周期"},"exit_":{"type":"range","min":5,"max":30,"desc":"退出周期"}}},
    "momentum":{"name":"动量追涨","desc":"强势追涨，弱势止损","params":{"ma_period":60,"mom_threshold":0.10},
        "param_meta":{"ma_period":{"type":"range","min":30,"max":120,"desc":"均线周期"},"mom_threshold":{"type":"range","min":0.05,"max":0.25,"step":0.01,"desc":"动量阈值"}}},
    "rsi_trailing":{"name":"RSI+移动止盈","desc":"RSI超卖买，移动止盈卖","params":{"oversold":25,"trail_pct":0.08,"rsi_exit":75},
        "param_meta":{"oversold":{"type":"range","min":15,"max":35,"desc":"超卖阈值"},"trail_pct":{"type":"range","min":0.03,"max":0.15,"step":0.01,"desc":"移动止盈%"},"rsi_exit":{"type":"range","min":60,"max":85,"desc":"RSI退出"}}},
}

# ═══════════ WebSocket 连接池 ═══════════
ws_clients = set()

@app.websocket("/ws/monitor")
async def ws_monitor(ws: WebSocket):
    await ws.accept()
    ws_clients.add(ws)
    try:
        while True:
            await asyncio.sleep(5)
            try:
                codes = ['sh600519','sh600036','sz000001','sz300750','sh601318','sh600276','sz002415','sh603259','sz000333','sz000858','sh601899','sz300059']
                rt = fetch_realtime(codes)
                # 简化返回
                alerts = []
                for raw_code, info in rt.items():
                    code = raw_code[2:]  # remove sh/sz prefix
                    # 拿历史算RSI
                    kl = fetch_kline(code, 100)
                    if kl:
                        closes=[k[2] for k in kl]
                        closes.append(info['price'])
                        r=calc_rsi(closes)
                    else: r=50
                    chg=(info['price']-info['yest_close'])/info['yest_close']*100
                    alerts.append({
                        'code':code,'name':info['name'],'price':info['price'],
                        'change':round(chg,2),'rsi':round(r,1),
                        'signal':'BUY' if r<25 else 'SELL' if r>80 else None
                    })
                await ws.send_json({'time':datetime.now().strftime('%H:%M:%S'),'data':alerts})
            except Exception:
                break
    except WebSocketDisconnect:
        pass
    finally:
        ws_clients.discard(ws)

# ═══════════ API ═══════════

@app.get("/api/stocks")
def api_stocks(): return STOCK_POOL

@app.get("/api/stock/{code}/history")
def api_history(code: str, days: int = 500):
    kl = fetch_kline(code, days)
    if not kl: return JSONResponse({"error":"无数据"},404)
    closes=[k[2] for k in kl]
    return {"code":code,"total_days":len(kl),
        "recent":[{"date":k[0],"open":k[1],"close":k[2],"high":k[3],"low":k[4],"volume":k[5]} for k in kl[-200:]],
        "last_close":closes[-1],"rsi":round(calc_rsi(closes),1),
        "ma20":round(calc_ma(closes,20),2),"ma60":round(calc_ma(closes,60),2),
        "change_20d":round((closes[-1]-closes[-20])/closes[-20]*100,2) if len(closes)>=20 else 0}


@app.get("/api/stock/{code}/detail")
def api_stock_detail(code: str):
    """单只股票完整详情"""
    kl = fetch_kline(code, 500)
    if not kl: return JSONResponse({"error":"无数据"}, 404)
    closes = [k[2] for k in kl]
    r = calc_rsi(closes)
    name = next((s["name"] for s in STOCK_POOL if s["code"]==code), code)

    # 计算多个指标
    def trend_angle(prices, n):
        if len(prices) < n: return 0
        y = prices[-n:]; x = list(range(n))
        n_ = n; sxy = sum(x[i]*y[i] for i in range(n_))
        sx = sum(x); sy = sum(y); sxx = sum(x[i]*x[i] for i in range(n_))
        slope = (n_*sxy - sx*sy) / (n_*sxx - sx*sx + 0.0001)
        return math.degrees(math.atan(slope))

    returns = [(closes[i]-closes[i-1])/closes[i-1] for i in range(1, len(closes))]
    vol = (sum(rr*rr for rr in returns)/len(returns))**0.5 * 100 if returns else 0

    # 回溯RSI序列(最近200天)
    rsi_series = [50]*14
    for i in range(1, len(closes)):
        g = [max(closes[j]-closes[j-1],0) for j in range(1,i+1)]
        ls = [max(closes[j-1]-closes[j],0) for j in range(1,i+1)]
        ag = sum(g[-14:])/14; al = sum(ls[-14:])/14
        rsi_series.append(100-100/(1+ag/(al+0.0001)))

    recent = [{"date": kl[i][0], "open": kl[i][1], "close": kl[i][2], "high": kl[i][3], "low": kl[i][4], "volume": kl[i][5], "rsi": round(rsi_series[i],1)} for i in range(max(0,len(kl)-200), len(kl))]

    ma20 = sum(closes[-20:])/20 if len(closes)>=20 else closes[-1]
    ma60 = sum(closes[-60:])/60 if len(closes)>=60 else closes[-1]

    return {
        "code": code, "name": name,
        "last_close": closes[-1], "rsi": round(r,1),
        "trend_60": round(trend_angle(closes, 60), 1),
        "trend_120": round(trend_angle(closes, 120), 1),
        "trend_252": round(trend_angle(closes, min(252, len(closes))), 1),
        "volatility": round(vol, 2),
        "ma20": round(ma20, 2), "ma60": round(ma60, 2),
        "high_60d": round(max(closes[-60:]), 2), "low_60d": round(min(closes[-60:]), 2),
        "recent": recent
    }

@app.get("/api/stock/{code}/realtime")
def api_realtime(code: str):
    rt_key = f'sh{code}' if code.startswith(('6','9')) else f'sz{code}'
    rt = fetch_realtime([rt_key])
    info = rt.get(rt_key)
    if not info: return JSONResponse({"error":"获取失败"},500)
    kl = fetch_kline(code, 100)
    if kl:
        closes=[k[2] for k in kl]; closes.append(info['price'])
        r=calc_rsi(closes)
    else: r=50
    return {**info,"rsi":round(r,1),"change_pct":round((info['price']-info['yest_close'])/info['yest_close']*100,2)}

@app.get("/api/dashboard")
def api_dashboard():
    codes = [s["code"] for s in STOCK_POOL[:30]]
    items = []; oversold = 0; overbought = 0
    for code in codes:
        kl = fetch_kline(code, 100)
        if not kl: continue
        closes = [k[2] for k in kl]
        name = next((s["name"] for s in STOCK_POOL if s["code"]==code), code)
        r = round(calc_rsi(closes), 1)
        chg = round((closes[-1]-closes[-20])/closes[-20]*100, 2) if len(closes)>=20 else 0
        if r < 25: oversold += 1
        if r > 80: overbought += 1
        items.append({"code":code,"name":name,"close":closes[-1],"rsi":r,"change_20d":chg,"ma20":round(calc_ma(closes,20),2),"signal":"BUY" if r<25 else "SELL" if r>80 else None})
    return {"items":items,"oversold":oversold,"overbought":overbought,"sentiment":"偏弱·机会" if oversold>0 else "偏强·谨慎" if overbought>0 else "中性","total":len(items)}

@app.get("/api/strategies")
def api_strategies():
    result = [{"id":k,"name":v["name"],"desc":v["desc"],"params":v["params"],"param_meta":v.get("param_meta",{})} for k,v in PRESET_STRATEGIES.items()]
    result.append({"id":"custom","name":"自定义策略","desc":"编写你自己的买卖条件","params":{"buy_condition":"rsi_14 < 25","sell_condition":"rsi_14 > 80"},"param_meta":{"buy_condition":{"type":"code","desc":"买入条件(Python表达式)"},"sell_condition":{"type":"code","desc":"卖出条件(Python表达式)"}}})
    return result

@app.post("/api/backtest")
def api_backtest(data: dict):
    codes = list(dict.fromkeys(data.get("codes",[])))
    sid = data.get("strategy","rsi")
    params = data.get("params",{})
    if sid=="custom" and "buy_condition" not in params:
        return JSONResponse({"error":"自定义策略需要 buy_condition 和 sell_condition"},400)
    if sid in PRESET_STRATEGIES:
        params = {**PRESET_STRATEGIES[sid]["params"], **params}
    else:
        params["strategy"] = "custom"
    results=[]
    total_pnl=0
    max_workers = min(32, max(1, len(codes)))
    with ThreadPoolExecutor(max_workers=max_workers) as pool:
        future_map = {pool.submit(run_backtest_detailed, code, {**params, "strategy":sid}): code for code in codes}
        for future in as_completed(future_map):
            code = future_map[future]
            try:
                r = future.result()
            except Exception:
                r = None
            if r:
                results.append({"code":code,**r})
                total_pnl += r["pnl"]
    results.sort(key=lambda x:x["pnl"], reverse=True)
    return {"strategy":PRESET_STRATEGIES.get(sid,{}).get("name","自定义"),"params":params,"results":results,
        "total_pnl":total_pnl,"total_trades":sum(r["trades"] for r in results),
        "avg_win_rate":round(sum(r["win_rate"] for r in results)/len(results),1) if results else 0,
        "avg_sharpe":round(sum(r["sharpe"] for r in results)/len(results),2) if results else 0,
        "max_dd":round(max(r["max_drawdown"] for r in results),2) if results else 0}

@app.post("/api/backtest/export")
def api_backtest_export(data: dict):
    """导出回测结果为 CSV"""
    codes = data.get("codes",[]); sid = data.get("strategy","rsi"); params = data.get("params",{})
    if sid in PRESET_STRATEGIES: params = {**PRESET_STRATEGIES[sid]["params"], **params}
    else: params["strategy"]="custom"
    output = io.StringIO(); w = csv.writer(output)
    w.writerow(["股票","买入日","买入价","卖出日","卖出价","盈亏","收益率%","持仓天数","最大浮盈%"])
    for code in codes:
        r = run_backtest_detailed(code, {**params, "strategy":sid})
        if r:
            for t in r["trade_list"]:
                w.writerow([code,t['buy_date'],t['buy_price'],t['sell_date'],t['sell_price'],t['pnl'],t['pct'],t['holding_days'],t['peak_dd']])
    output.seek(0)
    return StreamingResponse(iter([output.getvalue()]), media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=backtest_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"})

@app.post("/api/screen")
def api_screen(data: dict):
    started_at = time.time()
    codes = data.get("codes",[s["code"] for s in STOCK_POOL])
    name_map = {s["code"]: s["name"] for s in STOCK_POOL}

    def analyze_one(code):
        kl = fetch_kline(code,300)
        if not kl or len(kl)<60: return None
        closes=[k[2] for k in kl]
        def trend_angle(n):
            y=closes[-n:]; x=list(range(n))
            n_=n; sxy=sum(xi*yi for xi,yi in zip(x,y)); sx,sy=sum(x),sum(y); sxx=sum(xi*xi for xi in x)
            slope=(n_*sxy-sx*sy)/(n_*sxx-sx*sx+0.0001)
            return math.degrees(math.atan(slope))
        a60=trend_angle(60); a120=trend_angle(120)
        returns=[(closes[i]-closes[i-1])/closes[i-1] for i in range(1,len(closes))]
        vol=(sum(r*r for r in returns)/len(returns))**0.5*100
        r=calc_rsi(closes)
        score=a60*0.35+a120*0.25+max(0,50-abs(r-50))*0.2+max(0,50-abs(vol-2)*15)*0.2
        return {"code":code,"name":name_map.get(code,code),"score":round(score,1),"trend_60":round(a60,1),"trend_120":round(a120,1),"volatility":round(vol,2),"rsi":round(r,1),"last_close":closes[-1]}

    results=[]
    codes = list(dict.fromkeys(codes))
    max_workers = min(48, max(1, len(codes)))
    with ThreadPoolExecutor(max_workers=max_workers) as pool:
        for future in as_completed([pool.submit(analyze_one, code) for code in codes]):
            try:
                item = future.result()
            except Exception:
                item = None
            if item: results.append(item)
    results.sort(key=lambda x:x["score"],reverse=True)
    return {
        "results": results,
        "total": len(codes),
        "analyzed": len(results),
        "skipped": max(0, len(codes) - len(results)),
        "elapsed": round(time.time() - started_at, 2)
    }

@app.post("/api/analyze")
def api_analyze(data: dict):
    api_key = load_anthropic_api_key()
    if not api_key: return JSONResponse({"error":"需要 ANTHROPIC_API_KEY"},400)
    text = data.get("orders_text","")
    if not text: return JSONResponse({"error":"请提供回测数据"},400)
    import requests as req
    payload = {"model":"claude-sonnet-4-6","max_tokens":2048,"temperature":0.3,"messages":[{"role":"user","content":f"分析以下回测结果，输出JSON(只输出JSON): {{{{\"summary\":\"总体评价\",\"loss_patterns\":\"亏损共性\",\"win_patterns\":\"盈利共性\",\"suggestions\":\"优化建议\",\"next_steps\":\"下一步\"}}}}\n\n{text}"}]}
    try:
        r=req.post("https://api.anthropic.com/v1/messages",headers={"x-api-key":api_key,"anthropic-version":"2023-06-01","content-type":"application/json"},json=payload,timeout=45)
        r.raise_for_status()
        resp=r.json()["content"][0]["text"]
        m=re.search(r'\{.*\}',resp,re.DOTALL)
        return json.loads(m.group()) if m else {"raw":resp}
    except (WebSocketDisconnect, Exception):
        return JSONResponse({"error":str(e)},500)

@app.post("/api/report")
def api_report(data: dict):
    codes=data.get("codes",["600519","000001","300750"])
    rt_codes=[f'sh{c}' if c.startswith(('6','9')) else f'sz{c}' for c in codes]
    rt=fetch_realtime(rt_codes)
    lines=[f"# A股市场观察 {datetime.now().strftime('%Y-%m-%d')}\n"]
    for code in codes:
        rt_key=f'sh{code}' if code.startswith(('6','9')) else f'sz{code}'
        info=rt.get(rt_key)
        if not info: continue
        kl=fetch_kline(code,200)
        rsi_v=50; ma20=info['price']
        if kl:
            closes=[k[2] for k in kl]; closes.append(info['price'])
            rsi_v=round(calc_rsi(closes),1); ma20=calc_ma(closes,20)
        chg=(info['price']-info['yest_close'])/info['yest_close']*100
        sig="🔵超卖" if rsi_v<25 else "🔴超买" if rsi_v>80 else "🟡偏高" if rsi_v>70 else "⚪正常"
        lines.append(f"## {info['name']} ({code})")
        lines.append(f"- 现价:{info['price']:.2f} | 涨跌:{chg:+.2f}%")
        lines.append(f"- RSI(14):{rsi_v} {sig} | MA20:{ma20:.2f}")
        lines.append("")
    report="\n".join(lines)
    try:
        from docx import Document
        doc=Document(); doc.add_heading(f'A股观察 {datetime.now().strftime("%Y-%m-%d")}',0)
        for line in lines[1:]:
            if line.startswith('## '): doc.add_heading(line[3:],1)
            elif line.startswith('- '): doc.add_paragraph(line[2:])
        buf=io.BytesIO(); doc.save(buf); buf.seek(0)
        import base64
        return {"markdown":report,"docx_base64":base64.b64encode(buf.read()).decode()}
    except: return {"markdown":report}

app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/")
def index(): return FileResponse("static/index.html")

if __name__=="__main__":
    import uvicorn
    print("🟢 QuantDesk 启动 → http://localhost:8888")
    uvicorn.run(app, host="0.0.0.0", port=8888, log_level="warning")
