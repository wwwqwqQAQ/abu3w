#!/usr/bin/env python3
"""QuantDesk — 量化交易工作站后端"""
import subprocess, json, math, os, sys, io, csv, time, asyncio, hashlib, re, threading
from datetime import datetime, timedelta
from concurrent.futures import ThreadPoolExecutor, as_completed
from fastapi import FastAPI, Query, WebSocket, WebSocketDisconnect
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse, Response, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import warnings; warnings.filterwarnings('ignore')

sys.path.insert(0, os.path.dirname(__file__))
try:
    from predict import train as predict_train
    PREDICT_AVAILABLE = True
except ImportError:
    PREDICT_AVAILABLE = False

try:
    from radar import news_fetcher as radar_news
    from radar import sentiment as radar_sentiment
    from radar import risk as radar_risk
    from radar import scorer as radar_scorer
    RADAR_AVAILABLE = True
except ImportError:
    RADAR_AVAILABLE = False

try:
    from predict import walkforward as predict_wf
    WF_AVAILABLE = True
except ImportError:
    WF_AVAILABLE = False

sys.path.insert(0, os.path.dirname(__file__))
app = FastAPI(title="QuantDesk")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# ═══════════ 数据 ═══════════
KLINE_CACHE = {}
KLINE_CACHE_TTL = 60 * 60 * 24
SUPPORT_DIR = os.path.expanduser("~/Library/Application Support/QuantDesk")
KLINE_CACHE_DIR = os.path.join(SUPPORT_DIR, "cache/kline")
KLINE_CONNECT_TIMEOUT = "1"
KLINE_MAX_TIME = "3"
KLINE_PROCESS_TIMEOUT = 4
PREDICTION_MAX_STALE_DAYS = 10
HISTORY_WINDOW_DAYS = 1500
FORECAST_HISTORY_DAYS = 1500
PREDICTION_SIGNALS_TTL = 180
PREDICTION_SIGNALS_CACHE = {"key": None, "ts": 0, "data": None}
PREDICTION_SIGNALS_LOCK = threading.Lock()
PREDICTION_SIGNALS_PATH = os.path.join(SUPPORT_DIR, "cache/predict_signals.json")
UNIVERSE_CACHE_PATH = os.path.join(SUPPORT_DIR, "universe_all_a.json")
DECISION_CONTEXT_PATH = os.path.join(SUPPORT_DIR, "decision_context.json")
PREDICTION_JOB = {
    "running": False,
    "started_at": None,
    "finished_at": None,
    "error": None,
    "message": None,
}
PREDICTION_JOB_LOCK = threading.Lock()
WF_JOB = {
    "running": False,
    "started_at": None,
    "finished_at": None,
    "params": None,
    "result": None,
    "error": None,
    "message": None,
}
WF_JOB_LOCK = threading.Lock()


def cache_key(code, days):
    return f"{market_prefix(code)}{code}_{int(days)}"


def cache_path(code, days):
    return os.path.join(KLINE_CACHE_DIR, f"{cache_key(code, days)}.json")


def market_prefix(code):
    if code == "000300" or code.startswith(("6", "9")):
        return "sh"
    return "sz"


def read_kline_cache(code, days, allow_stale=False):
    now = time.time()
    for cached_days in sorted({int(days), 500, 300, 200, 100}, reverse=True):
        if cached_days < int(days):
            continue
        key = cache_key(code, cached_days)
        item = KLINE_CACHE.get(key)
        if item and (allow_stale or now - item["time"] <= KLINE_CACHE_TTL):
            return item["data"][-int(days):]

        path = cache_path(code, cached_days)
        try:
            if not os.path.exists(path):
                continue
            if not allow_stale and now - os.path.getmtime(path) > KLINE_CACHE_TTL:
                continue
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            KLINE_CACHE[key] = {"time": os.path.getmtime(path), "data": data}
            return data[-int(days):]
        except Exception:
            continue
    return None


def write_kline_cache(code, days, data):
    if not data:
        return
    key = cache_key(code, days)
    KLINE_CACHE[key] = {"time": time.time(), "data": data}
    try:
        os.makedirs(KLINE_CACHE_DIR, exist_ok=True)
        with open(cache_path(code, days), "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, separators=(",", ":"))
    except Exception:
        pass


def load_anthropic_api_key():
    key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if key:
        return key

    support_dir = os.path.expanduser("~/Library/Application Support/QuantDesk")
    json_paths = [
        os.path.join(support_dir, "config.json"),
        os.path.expanduser("~/.quantdesk.json"),
    ]
    env_paths = [
        os.path.join(support_dir, ".env"),
        os.path.join(support_dir, "anthropic_api_key.txt"),
        os.path.expanduser("~/.quantdesk.env"),
    ]

    for path in json_paths:
        try:
            if not os.path.exists(path):
                continue
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            key = str(data.get("ANTHROPIC_API_KEY") or data.get("anthropic_api_key") or "").strip()
            if key:
                return key
        except Exception:
            continue

    for path in env_paths:
        try:
            if not os.path.exists(path):
                continue
            with open(path, "r", encoding="utf-8") as f:
                text = f.read().strip()
            if path.endswith(".txt") and "=" not in text:
                return text.splitlines()[0].strip()
            for line in text.splitlines():
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                name, value = line.split("=", 1)
                if name.strip() == "ANTHROPIC_API_KEY":
                    return value.strip().strip("\"'")
        except Exception:
            continue

    return ""


def fetch_kline(code, days=500):
    cached = read_kline_cache(code, days)
    if cached is not None:
        return cached

    m = market_prefix(code)
    url = f'http://money.finance.sina.com.cn/quotes_service/api/json_v2.php/CN_MarketData.getKLineData?symbol={m}{code}&scale=240&ma=no&datalen={days}'
    try:
        r = subprocess.run(
            ['curl','-sS','--connect-timeout',KLINE_CONNECT_TIMEOUT,'--max-time',KLINE_MAX_TIME,'--noproxy','*',url],
            capture_output=True,
            text=True,
            timeout=KLINE_PROCESS_TIMEOUT
        )
        if r.returncode != 0 or not r.stdout or r.stdout.startswith('null'):
            return read_kline_cache(code, days, allow_stale=True) or []
        data = [(k['day'],float(k['open']),float(k['close']),float(k['high']),float(k['low']),int(float(k['volume']))) for k in json.loads(r.stdout)]
        write_kline_cache(code, days, data)
        return data
    except Exception:
        return read_kline_cache(code, days, allow_stale=True) or []


def kline_age_days(kl):
    if not kl:
        return None
    try:
        last_date = datetime.strptime(kl[-1][0], "%Y-%m-%d")
        return (datetime.now() - last_date).days
    except (ValueError, TypeError, IndexError):
        return None


def is_kline_recent(kl, max_age_days=PREDICTION_MAX_STALE_DAYS):
    age = kline_age_days(kl)
    return age is not None and age <= max_age_days


def fetch_kline_prediction(code, days=500):
    data = read_kline_cache(code, days, allow_stale=True)
    if data and is_kline_recent(data):
        return data
    if code in ("000300", "000905", "000852", "000001"):
        data = fetch_kline(code, days)
        if data and is_kline_recent(data):
            return data
    return []

def fetch_realtime(codes):
    if not codes: return {}
    ids = ','.join(codes)
    try:
        r = subprocess.run(
            ['curl','-sS','--connect-timeout','1','--max-time','3','--noproxy','*','-H','Referer: https://finance.sina.com.cn','-H','User-Agent: Mozilla/5.0',f'http://hq.sinajs.cn/list={ids}'],
            capture_output=True,
            timeout=4
        )
        if r.returncode != 0:
            return {}
        text = r.stdout.decode('gbk', errors='replace')
    except Exception:
        return {}
    results = {}
    for line in text.strip().split('\n'):
        if not line or '=' not in line: continue
        c = line.split('=')[0].split('_')[-1]
        raw = line.split('"')[1] if '"' in line else ''
        if raw:
            p = raw.split(',')
            if len(p)>=6:
                try: results[c] = {'name':p[0],'open':float(p[1]),'yest_close':float(p[2]),'price':float(p[3]),'high':float(p[4]),'low':float(p[5]),'volume':0}
                except (ValueError, IndexError, TypeError):
                    pass
    return results

# ═══════════ 指标 ═══════════
def calc_rsi(closes, n=14):
    if len(closes)<n+1: return 50
    g=[max(closes[i]-closes[i-1],0) for i in range(1,len(closes))]
    l=[max(closes[i-1]-closes[i],0) for i in range(1,len(closes))]
    ag=sum(g[-n:])/n; al=sum(l[-n:])/n
    if al==0: return 100
    return 100-100/(1+ag/(al+0.0001))

def calc_ma(prices,n):
    if len(prices)<n: return sum(prices)/len(prices)
    return sum(prices[-n:])/n

def calc_adx(highs, lows, closes, n=14):
    """ADX — 趋势强度指标。>25 趋势市，<20 震荡市。"""
    if len(closes) < n + 1:
        return 20
    tr = []; plus_dm = []; minus_dm = []
    for i in range(1, len(closes)):
        hl = highs[i] - lows[i]
        hc = abs(highs[i] - closes[i-1])
        lc = abs(lows[i] - closes[i-1])
        tr.append(max(hl, hc, lc))
        up = highs[i] - highs[i-1]
        dn = lows[i-1] - lows[i]
        plus_dm.append(up if up > dn and up > 0 else 0)
        minus_dm.append(dn if dn > up and dn > 0 else 0)
    atr = sum(tr[:n]) / n
    pdi = sum(plus_dm[:n]) / n / (atr + 0.0001) * 100
    mdi = sum(minus_dm[:n]) / n / (atr + 0.0001) * 100
    dx = abs(pdi - mdi) / (pdi + mdi + 0.0001) * 100
    adx = dx
    for i in range(n, len(tr)):
        atr = (atr * (n - 1) + tr[i]) / n
        pdi = (pdi * (n - 1) + (plus_dm[i] / (atr + 0.0001) * 100)) / n
        mdi = (mdi * (n - 1) + (minus_dm[i] / (atr + 0.0001) * 100)) / n
        dx = abs(pdi - mdi) / (pdi + mdi + 0.0001) * 100
        adx = (adx * (n - 1) + dx) / n
    return round(adx, 1)

def calc_atr(highs, lows, closes, n=14):
    """ATR — 平均真实波幅，用于动态止损和仓位管理。"""
    if len(closes) < 2:
        return 0
    trs = []
    for i in range(1, len(closes)):
        hl = highs[i] - lows[i]
        hc = abs(highs[i] - closes[i-1])
        lc = abs(lows[i] - closes[i-1])
        trs.append(max(hl, hc, lc))
    if len(trs) < n:
        return sum(trs) / len(trs)
    atr = sum(trs[:n]) / n
    for i in range(n, len(trs)):
        atr = (atr * (n - 1) + trs[i]) / n
    return round(atr, 2)


def _file_mtime(path):
    try:
        return os.path.getmtime(path)
    except Exception:
        return 0


def _prediction_signals_cache_key(codes):
    model_path = getattr(predict_train, "MODEL_PATH", "") if PREDICT_AVAILABLE else ""
    metrics_path = getattr(predict_train, "METRICS_PATH", "") if PREDICT_AVAILABLE else ""
    raw = "|".join([
        ",".join(codes),
        str(PREDICTION_MAX_STALE_DAYS),
        str(_file_mtime(model_path)),
        str(_file_mtime(metrics_path)),
    ])
    return hashlib.sha1(raw.encode("utf-8")).hexdigest()


def _read_json_file(path):
    try:
        if not os.path.exists(path):
            return None
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def _write_json_file(path, data):
    try:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        tmp_path = path + ".tmp"
        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, separators=(",", ":"))
        os.replace(tmp_path, path)
    except Exception:
        pass


def _load_prediction_snapshot(cache_key):
    data = _read_json_file(PREDICTION_SIGNALS_PATH)
    if not data or data.get("_cache_key") != cache_key:
        return None
    return data


def _prediction_job_snapshot():
    with PREDICTION_JOB_LOCK:
        return dict(PREDICTION_JOB)


def _build_prediction_signals(codes, name_map, cache_key):
    model = predict_train.load_model()
    if model is None:
        return {"error": "模型未训练，请先 POST /api/predict/train"}
    results = predict_train.predict_all(fetch_kline_prediction, codes, name_map, benchmark_code="000300", model=model)
    results = [apply_decision_overlay(item) for item in results]
    status = predict_train.get_status()
    return {
        "_cache_key": cache_key,
        "predictions": results,
        "metrics": status,
        "updated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "total": len(results),
        "cache": {"hit": False, "ttl_seconds": PREDICTION_SIGNALS_TTL, "source": "fresh_full"},
        "data_quality": {
            "source": "local_cache_recent_only",
            "max_stale_days": PREDICTION_MAX_STALE_DAYS,
            "skipped": max(0, len(codes) - len(results)),
        },
        "prediction_job": _prediction_job_snapshot(),
    }


def _run_prediction_job(cache_key, codes, name_map):
    with PREDICTION_JOB_LOCK:
        if PREDICTION_JOB.get("running"):
            return
        PREDICTION_JOB.update({
            "running": True,
            "started_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "finished_at": None,
            "error": None,
            "message": "后台刷新ML预测",
        })
    try:
        data = _build_prediction_signals(codes, name_map, cache_key)
        if data.get("error"):
            raise RuntimeError(data["error"])
        with PREDICTION_SIGNALS_LOCK:
            PREDICTION_SIGNALS_CACHE.update({"key": cache_key, "ts": time.time(), "data": data})
        _write_json_file(PREDICTION_SIGNALS_PATH, data)
        with PREDICTION_JOB_LOCK:
            PREDICTION_JOB.update({
                "running": False,
                "finished_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "error": None,
                "message": "ML预测已刷新",
            })
    except Exception as e:
        with PREDICTION_JOB_LOCK:
            PREDICTION_JOB.update({
                "running": False,
                "finished_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "error": str(e),
                "message": "ML预测刷新失败",
            })


def _trend_angle_from_closes(closes, n):
    if len(closes) <= n or closes[-n] <= 0:
        return 0.0
    pct_move = (closes[-1] - closes[-n]) / closes[-n] * 100
    return round(math.degrees(math.atan(pct_move / n)), 2)


def _volatility_from_closes(closes, n=20):
    if len(closes) <= n:
        return 0.02
    rets = []
    for i in range(len(closes) - n + 1, len(closes)):
        if closes[i - 1] > 0:
            rets.append((closes[i] - closes[i - 1]) / closes[i - 1])
    if len(rets) < 2:
        return 0.02
    avg = sum(rets) / len(rets)
    var = sum((r - avg) ** 2 for r in rets) / (len(rets) - 1)
    return round(math.sqrt(var), 4)


def _atr_pct_from_kline(kl, n=14):
    if len(kl) <= n:
        return 0.0
    trs = []
    start = max(1, len(kl) - n)
    for i in range(start, len(kl)):
        high = kl[i][3]
        low = kl[i][4]
        prev_close = kl[i - 1][2]
        trs.append(max(high - low, abs(high - prev_close), abs(low - prev_close)))
    close = kl[-1][2]
    return round((sum(trs) / len(trs)) / close * 100, 3) if close > 0 and trs else 0.0


def risk_kline_features(code, kl=None, ml_pred=None):
    features = {
        "rsi_14": 50,
        "trend_20": 0,
        "trend_60": 0,
        "volatility_20": 0.02,
        "ma20_dev": 0,
        "ma60_dev": 0,
        "price_position": 0.5,
        "atr_14_pct": 0,
    }
    model_features = (ml_pred or {}).get("features_used", {}) or {}
    for key in features:
        if model_features.get(key) is not None:
            features[key] = model_features.get(key)

    if not kl or len(kl) < 20:
        return features
    closes = [k[2] for k in kl]
    highs = [k[3] for k in kl]
    lows = [k[4] for k in kl]
    features["rsi_14"] = round(calc_rsi(closes), 1)
    features["volatility_20"] = _volatility_from_closes(closes, 20)
    features["trend_20"] = _trend_angle_from_closes(closes, 20)
    features["atr_14_pct"] = _atr_pct_from_kline(kl, 14)
    ma20 = calc_ma(closes, 20)
    features["ma20_dev"] = round((closes[-1] - ma20) / ma20 * 100, 2) if ma20 > 0 else 0
    if len(closes) >= 60:
        features["trend_60"] = _trend_angle_from_closes(closes, 60)
        ma60 = calc_ma(closes, 60)
        features["ma60_dev"] = round((closes[-1] - ma60) / ma60 * 100, 2) if ma60 > 0 else 0
        h60 = max(highs[-60:])
        l60 = min(lows[-60:])
        features["price_position"] = round((closes[-1] - l60) / (h60 - l60), 3) if h60 > l60 else 0.5
    return features


def _wf_job_snapshot():
    with WF_JOB_LOCK:
        return dict(WF_JOB)


def _run_walkforward_job(params):
    def progress(message):
        with WF_JOB_LOCK:
            WF_JOB["message"] = message

    with WF_JOB_LOCK:
        WF_JOB.update({
            "running": True,
            "started_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "finished_at": None,
            "params": params,
            "result": None,
            "error": None,
            "message": "启动策略验证",
        })
    try:
        result = predict_wf.run_walkforward(fetch_kline, progress_cb=progress, **params)
        with WF_JOB_LOCK:
            WF_JOB.update({
                "running": False,
                "finished_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "result": result,
                "error": result.get("error") if isinstance(result, dict) else None,
                "message": "策略验证完成",
            })
    except Exception as e:
        with WF_JOB_LOCK:
            WF_JOB.update({
                "running": False,
                "finished_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "error": str(e),
                "message": "策略验证失败",
            })

# ═══════════ 增强回测 ═══════════
def run_backtest_detailed(code, params):
    """详细回测 — 返回每笔交易明细 + 风险指标"""
    kl = fetch_kline(code, 500)
    if not kl: return None

    o=[k[1] for k in kl]; c=[k[2] for k in kl]; h=[k[3] for k in kl]; l=[k[4] for k in kl]; v=[k[5] for k in kl]
    trades=[]; holding=False; bp=0; bd=''; cur=100000; max_price=0; max_dd=0; peak=100000
    # 风控参数
    risk_per_trade = float(params.get("risk_pct", 0.02))  # 单笔风险暴露
    max_portfolio_dd = float(params.get("max_portfolio_dd", 0.25))  # 组合最大回撤停止线
    trading_stopped = False

    sid = params.get("strategy","rsi")
    # 构建 RSI — 预计算涨跌幅避免 O(n²) 重复遍历
    r=[50]*14
    gains=[max(c[i]-c[i-1],0) for i in range(1,len(c))]
    losses=[max(c[i-1]-c[i],0) for i in range(1,len(c))]
    for i in range(len(gains)):
        w=min(14,i+1)
        ag=sum(gains[max(0,i-13):i+1])/w
        al=sum(losses[max(0,i-13):i+1])/w
        r.append(100-100/(1+ag/(al+0.0001)))

    # 计算均线
    def ma_at(prices,n,i):
        if i<n: return sum(prices[:i+1])/(i+1)
        return sum(prices[i-n+1:i+1])/n

    # 趋势角度（线性回归斜率→角度）
    def trend_at(prices, n, i):
        if i < n: return 0
        seg = prices[i-n+1:i+1]
        xs = list(range(n))
        sx = sum(xs); sy = sum(seg)
        sxx = sum(x*x for x in xs); sxy = sum(xs[j]*seg[j] for j in range(n))
        denom = n*sxx - sx*sx
        if abs(denom) < 0.0001: return 0
        return math.degrees(math.atan((n*sxy - sx*sy) / denom))

    # 量比
    def vol_ratio_at(vols, n, i):
        if i < n: return 1.0
        avg = sum(vols[i-n+1:i+1]) / n
        return vols[i] / avg if avg > 0 else 1.0

    # ATR for position sizing
    def atr_at(i, n=14):
        if i < 2: return (h[i]-l[i]) * 0.5
        trs = []
        start = max(1, i-n+1)
        for j in range(start, i+1):
            hl = h[j] - l[j]
            hc = abs(h[j] - c[j-1])
            lc = abs(l[j] - c[j-1])
            trs.append(max(hl, hc, lc))
        return sum(trs) / len(trs)

    def calc_shares(buy_price, atr_val):
        """ATR-based position sizing: risk a fixed % of capital per trade, capped at available capital."""
        if atr_val <= 0 or buy_price <= 0:
            return max(100, int(cur / buy_price / 100) * 100)
        risk_amount = cur * risk_per_trade
        stop_dist = atr_val * 2.0
        raw_shares = int(risk_amount / stop_dist / 100) * 100
        max_shares = int(cur / buy_price / 100) * 100  # 不能超过可用资金
        return max(100, min(raw_shares, max_shares))

    for i in range(60,len(kl)-1):
        # 更新峰值和回撤
        peak = max(peak, cur)
        dd = (peak-cur)/peak if peak>0 else 0
        max_dd = max(max_dd, dd)
        # 组合最大回撤停止：亏太多就停
        if dd >= max_portfolio_dd:
            trading_stopped = True
        if trading_stopped:
            if holding:
                # 清仓离场
                sp = o[i+1]; sd = kl[i+1][0]
                sh = int(cur/bp/100)*100
                if sh >= 100:
                    days = i - next((j for j in range(len(kl)) if kl[j][0]==bd), i)
                    pnl = (sp-bp)*sh; cur += pnl
                    trades.append({
                        'buy_date':bd,'buy_price':round(bp,2),'sell_date':sd,'sell_price':round(sp,2),
                        'pnl':round(pnl,2),'pct':round((sp-bp)/bp*100,2),
                        'holding_days':days,'peak_dd':round((max_price-bp)/bp*100,2),
                        'exit_reason':'组合回撤止损'
                    })
                holding=False
            continue

        if not holding:
            should_buy = False

            if sid=="rsi":
                os=params.get("oversold",25); should_buy=r[i]<os
            elif sid=="ma_cross":
                f=params.get("fast",5); s=params.get("slow",20)
                mf=ma_at(c,f,i); mf_prev=ma_at(c,f,i-1)
                ms=ma_at(c,s,i); ms_prev=ma_at(c,s,i-1)
                should_buy = mf_prev<=ms_prev and mf>ms
            elif sid=="turtle":
                e=params.get("entry",20); should_buy=h[i]>=max(h[max(0,i-e):i])
            elif sid=="momentum":
                mp=params.get("ma_period",60); mt=params.get("mom_threshold",0.10)
                mom=(c[i]-c[i-20])/c[i-20] if i>=20 else 0
                should_buy = c[i]>ma_at(c,mp,i) and mom>mt and r[i]>60
            elif sid=="rsi_trailing":
                os=params.get("oversold",25); should_buy=r[i]<os
            elif sid=="consensus":
                # 三策略共识：RSI + MA Cross + Momentum，≥2 票买入
                os_c = params.get("oversold", 25)
                f_c = params.get("fast", 5); s_c = params.get("slow", 20)
                mp_c = params.get("ma_period", 60); mt_c = params.get("mom_threshold", 0.10)
                mom_c = (c[i]-c[i-20])/c[i-20] if i>=20 else 0
                mf_c = ma_at(c,f_c,i); mf_prev_c = ma_at(c,f_c,i-1)
                ms_c = ma_at(c,s_c,i); ms_prev_c = ma_at(c,s_c,i-1)
                votes = 0
                if r[i] < os_c: votes += 1
                if mf_prev_c <= ms_prev_c and mf_c > ms_c: votes += 1
                if c[i] > ma_at(c,mp_c,i) and mom_c > mt_c and r[i] > 55: votes += 1
                should_buy = votes >= params.get("min_votes", 2)
            elif sid=="adaptive":
                # 自适应策略：ADX判市 → 趋势追涨 / 震荡抄底 / 过渡等金叉
                adx_val = calc_adx(h, l, c, 14)
                adx_t = params.get("adx_trend", 25)
                adx_c = params.get("adx_choppy", 15)
                rsi_os = params.get("rsi_os", 22)
                trend_20_adaptive = trend_at(c, 20, i)
                vol_ok_adaptive = vol_ratio_at(vols, 20, i) > 0.7

                if adx_val > adx_t:
                    # 趋势市：追涨 — 价格在MA20之上 + 趋势向上 + 放量
                    should_buy = c[i] > ma_at(c, 20, i) and trend_20_adaptive > 3 and vol_ok_adaptive
                elif adx_val < adx_c:
                    # 震荡市：抄底 — RSI超卖 + 不破MA60
                    should_buy = r[i] < rsi_os and c[i] > ma_at(c, 60, i) * 0.92
                else:
                    # 过渡市：金叉 + 放量
                    mf_ad = ma_at(c, 5, i); mf_prev_ad = ma_at(c, 5, i-1)
                    ms_ad = ma_at(c, 20, i); ms_prev_ad = ma_at(c, 20, i-1)
                    should_buy = mf_prev_ad <= ms_prev_ad and mf_ad > ms_ad and vol_ok_adaptive
            elif sid=="custom":
                # 自定义表达式求值
                should_buy = eval_condition(params.get("buy_condition","False"), c,o,h,l,r,i)

            if should_buy:
                bp=o[i+1]; bd=kl[i+1][0]; max_price=bp; holding=True

        elif holding:
            should_sell = False
            max_price = max(max_price, h[i])

            if sid=="rsi":
                ob=params.get("overbought",80); should_sell=r[i]>ob
            elif sid=="ma_cross":
                f=params.get("fast",5); s=params.get("slow",20)
                should_sell = ma_at(c,f,i)<=ma_at(c,s,i)
            elif sid=="turtle":
                x=params.get("exit_",10); should_sell=l[i]<=min(l[max(0,i-x):i])
            elif sid=="momentum":
                should_sell = c[i]<ma_at(c,20,i) or r[i]<40
            elif sid=="rsi_trailing":
                tp=params.get("trail_pct",0.08); re=params.get("rsi_exit",75)
                should_sell = c[i]<=max_price*(1-tp) or r[i]>re
            elif sid=="consensus":
                # 任一子策略说卖就卖，或 ATR 硬止损
                ob_c = params.get("overbought", 80)
                f_c = params.get("fast", 5); s_c = params.get("slow", 20)
                atr_val = atr_at(i); atr_pct = atr_val / c[i] if c[i] > 0 else 0.03
                hard_stop = c[i] <= max_price * (1 - atr_pct * 3)
                should_sell = r[i] > ob_c or ma_at(c,f_c,i) <= ma_at(c,s_c,i) or c[i] < ma_at(c,20,i) or hard_stop
            elif sid=="adaptive":
                # 自适应卖出：ATR止损 + RSI超买 + 趋势反转
                rsi_ob = params.get("rsi_ob", 78)
                atr_val = atr_at(i); atr_pct = atr_val / c[i] if c[i] > 0 else 0.03
                # ATR动态止损（2x ATR回撤就卖）
                hard_stop = c[i] <= max_price * (1 - atr_pct * 2)
                # 趋势反转（20日趋势变负）
                trend_reverse = trend_at(c, 20, i) < -5
                # RSI超买
                rsi_over = r[i] > rsi_ob
                # 跌破MA20且放量下跌
                breakdown = c[i] < ma_at(c, 20, i) and vol_ratio_at(vols, 20, i) > 1.2
                should_sell = hard_stop or trend_reverse or rsi_over or breakdown
            elif sid=="custom":
                should_sell = eval_condition(params.get("sell_condition","False"), c,o,h,l,r,i)

            if should_sell:
                sp = o[i+1]; sd = kl[i+1][0]
                atr_val = atr_at(i)
                sh = calc_shares(bp, atr_val)
                if sh>=100:
                    days = i - next((j for j in range(len(kl)) if kl[j][0]==bd), i)
                    # A股真实成本: 印花税0.05%(卖) + 佣金0.025%×2 + 过户费0.002% + 滑点0.03%×2
                    cost_rate = 0.0005 + 0.00025*2 + 0.00002 + 0.0003*2
                    gross_pnl = (sp-bp)*sh
                    cost = (bp + sp) * sh * cost_rate
                    pnl = gross_pnl - cost; cur += pnl
                    trades.append({
                        'buy_date':bd,'buy_price':round(bp,2),'sell_date':sd,'sell_price':round(sp,2),
                        'pnl':round(pnl,2),'pct':round((sp-bp)/bp*100 - cost_rate*2, 2),
                        'holding_days':days,'peak_dd':round((max_price-bp)/bp*100,2)
                    })
                holding=False

    if holding and not trading_stopped:
        sp=c[-1]; sd=kl[-1][0]
        atr_val = atr_at(len(kl)-2)
        sh = calc_shares(bp, atr_val)
        if sh>=100:
            days=len(kl)-1-next((j for j in range(len(kl)) if kl[j][0]==bd),0)
            cost_rate = 0.0005 + 0.00025*2 + 0.00002 + 0.0003*2
            gross_pnl = (sp-bp)*sh
            cost = (bp + sp) * sh * cost_rate
            pnl = gross_pnl - cost; cur += pnl
            trades.append({
                'buy_date':bd,'buy_price':round(bp,2),'sell_date':sd,'sell_price':round(sp,2),
                'pnl':round(pnl,2),'pct':round((sp-bp)/bp*100 - cost_rate*2, 2),
                'holding_days':days,'peak_dd':round((max_price-bp)/bp*100,2)
            })

    total_pnl=cur-100000
    wins=sum(1 for t in trades if t['pnl']>0)
    total=len(trades)

    # Sharpe ratio (年化) — 至少5笔交易才有统计意义
    if total>=5:
        returns=[t['pct']/100 for t in trades]
        avg_ret=sum(returns)/len(returns)
        std_ret=(sum((r-avg_ret)**2 for r in returns)/len(returns))**0.5
        sharpe = round(avg_ret/std_ret*(total**0.5),2) if std_ret>0 else 0
    else:
        sharpe=0

    # 盈亏比
    avg_win=sum(t['pct'] for t in trades if t['pnl']>0)/max(1,sum(1 for t in trades if t['pnl']>0))
    avg_loss=sum(t['pct'] for t in trades if t['pnl']<0)/max(1,sum(1 for t in trades if t['pnl']<0))

    return {
        'pnl':total_pnl,'trades':total,'wins':wins,
        'win_rate':round(wins/total*100,1) if total else 0,
        'avg_pct':round(sum(t['pct'] for t in trades)/total,2) if total else 0,
        'sharpe':sharpe,'max_drawdown':round(max_dd*100,2),
        'avg_win':round(avg_win,2),'avg_loss':round(avg_loss,2),
        'profit_factor':round(abs(avg_win/avg_loss),1) if avg_loss!=0 else 99,
        'trade_list':trades,
    }

def eval_condition(expr, c, o, h, l, rsi_vals, i):
    """安全求值用户自定义条件表达式"""
    if not expr or expr.strip()=='False': return False
    try:
        local_vars = {
            'close': c[i], 'open': o[i], 'high': h[i], 'low': l[i],
            'rsi_14': rsi_vals[i],
            'prev_close': c[i-1] if i>0 else c[i],
            'ma_20': sum(c[max(0,i-19):i+1])/min(20,i+1),
            'ma_60': sum(c[max(0,i-59):i+1])/min(60,i+1),
            'ma_120': sum(c[max(0,i-119):i+1])/min(120,i+1),
            'change_5d': (c[i]-c[i-5])/c[i-5]*100 if i>=5 else 0,
            'change_20d': (c[i]-c[i-20])/c[i-20]*100 if i>=20 else 0,
            'vol_20': (sum((c[j]-sum(c[max(0,i-19):i+1])/min(20,i+1))**2 for j in range(max(0,i-19),i+1))/min(20,i+1))**0.5/c[i]*100 if i>=19 else 1,
            'math': math,
        }
        return bool(eval(expr, {"__builtins__": {}}, local_vars))
    except Exception:
        return False

# ═══════════ 股票池 ═══════════
STOCK_POOL = [
    {"code":"600519","name":"茅台","market":"sh"},{"code":"600036","name":"招行","market":"sh"},
    {"code":"601318","name":"中国平安","market":"sh"},{"code":"600276","name":"恒瑞医药","market":"sh"},
    {"code":"600900","name":"长江电力","market":"sh"},{"code":"601166","name":"兴业银行","market":"sh"},
    {"code":"600030","name":"中信证券","market":"sh"},{"code":"600887","name":"伊利股份","market":"sh"},
    {"code":"601398","name":"工商银行","market":"sh"},{"code":"601899","name":"紫金矿业","market":"sh"},
    {"code":"603259","name":"药明康德","market":"sh"},{"code":"601012","name":"隆基绿能","market":"sh"},
    {"code":"601888","name":"中国中免","market":"sh"},{"code":"600809","name":"山西汾酒","market":"sh"},
    {"code":"601857","name":"中国石油","market":"sh"},{"code":"600028","name":"中国石化","market":"sh"},
    {"code":"601288","name":"农业银行","market":"sh"},{"code":"600690","name":"海尔智家","market":"sh"},
    {"code":"000001","name":"平安银行","market":"sz"},{"code":"000858","name":"五粮液","market":"sz"},
    {"code":"002415","name":"海康威视","market":"sz"},{"code":"300750","name":"宁德时代","market":"sz"},
    {"code":"000333","name":"美的集团","market":"sz"},{"code":"002594","name":"比亚迪","market":"sz"},
    {"code":"000568","name":"泸州老窖","market":"sz"},{"code":"300059","name":"东方财富","market":"sz"},
    {"code":"002475","name":"立讯精密","market":"sz"},{"code":"000725","name":"京东方","market":"sz"},
    {"code":"002714","name":"牧原股份","market":"sz"},{"code":"300015","name":"爱尔眼科","market":"sz"},
    {"code":"000063","name":"中兴通讯","market":"sz"},{"code":"002352","name":"顺丰控股","market":"sz"},
    {"code":"000002","name":"万科A","market":"sz"},{"code":"002304","name":"洋河股份","market":"sz"},
    {"code":"601668","name":"中国建筑","market":"sh"},{"code":"601988","name":"中国银行","market":"sh"},
    {"code":"601328","name":"交通银行","market":"sh"},{"code":"600000","name":"浦发银行","market":"sh"},
    {"code":"600048","name":"保利发展","market":"sh"},{"code":"600050","name":"中国联通","market":"sh"},
    {"code":"600104","name":"上汽集团","market":"sh"},{"code":"600111","name":"北方稀土","market":"sh"},
    {"code":"600150","name":"中国船舶","market":"sh"},{"code":"600196","name":"复星医药","market":"sh"},
    {"code":"600309","name":"万华化学","market":"sh"},{"code":"600406","name":"国电南瑞","market":"sh"},
    {"code":"600436","name":"片仔癀","market":"sh"},{"code":"600438","name":"通威股份","market":"sh"},
    {"code":"600489","name":"中金黄金","market":"sh"},{"code":"600547","name":"山东黄金","market":"sh"},
    {"code":"600570","name":"恒生电子","market":"sh"},{"code":"600585","name":"海螺水泥","market":"sh"},
    {"code":"600660","name":"福耀玻璃","market":"sh"},{"code":"600674","name":"川投能源","market":"sh"},
    {"code":"600745","name":"闻泰科技","market":"sh"},{"code":"600760","name":"中航沈飞","market":"sh"},
    {"code":"600795","name":"国电电力","market":"sh"},{"code":"600837","name":"海通证券","market":"sh"},
    {"code":"600845","name":"宝信软件","market":"sh"},{"code":"600893","name":"航发动力","market":"sh"},
    {"code":"600905","name":"三峡能源","market":"sh"},{"code":"600941","name":"中国移动","market":"sh"},
    {"code":"600958","name":"东方证券","market":"sh"},{"code":"600999","name":"招商证券","market":"sh"},
    {"code":"601006","name":"大秦铁路","market":"sh"},{"code":"601009","name":"南京银行","market":"sh"},
    {"code":"601088","name":"中国神华","market":"sh"},{"code":"601100","name":"恒立液压","market":"sh"},
    {"code":"601111","name":"中国国航","market":"sh"},{"code":"601138","name":"工业富联","market":"sh"},
    {"code":"601155","name":"新城控股","market":"sh"},{"code":"601186","name":"中国铁建","market":"sh"},
    {"code":"601211","name":"国泰君安","market":"sh"},{"code":"601225","name":"陕西煤业","market":"sh"},
    {"code":"601229","name":"上海银行","market":"sh"},{"code":"601997","name":"贵阳银行","market":"sh"},
    {"code":"601319","name":"中国人保","market":"sh"},{"code":"601336","name":"新华保险","market":"sh"},
    {"code":"601360","name":"三六零","market":"sh"},{"code":"601390","name":"中国中铁","market":"sh"},
    {"code":"601601","name":"中国太保","market":"sh"},{"code":"601628","name":"中国人寿","market":"sh"},
    {"code":"601633","name":"长城汽车","market":"sh"},{"code":"601688","name":"华泰证券","market":"sh"},
    {"code":"601728","name":"中国电信","market":"sh"},{"code":"601766","name":"中国中车","market":"sh"},
    {"code":"601800","name":"中国交建","market":"sh"},{"code":"601816","name":"京沪高铁","market":"sh"},
    {"code":"601818","name":"光大银行","market":"sh"},{"code":"601838","name":"成都银行","market":"sh"},
    {"code":"601868","name":"中国能建","market":"sh"},{"code":"601872","name":"招商轮船","market":"sh"},
    {"code":"601919","name":"中远海控","market":"sh"},{"code":"601939","name":"建设银行","market":"sh"},
    {"code":"601985","name":"中国核电","market":"sh"},{"code":"601989","name":"中国重工","market":"sh"},
    {"code":"603288","name":"海天味业","market":"sh"},{"code":"603501","name":"韦尔股份","market":"sh"},
    {"code":"603799","name":"华友钴业","market":"sh"},{"code":"603986","name":"兆易创新","market":"sh"},
    {"code":"688008","name":"澜起科技","market":"sh"},{"code":"688012","name":"中微公司","market":"sh"},
    {"code":"688036","name":"传音控股","market":"sh"},{"code":"688111","name":"金山办公","market":"sh"},
    {"code":"688126","name":"沪硅产业","market":"sh"},{"code":"688169","name":"石头科技","market":"sh"},
    {"code":"688187","name":"时代电气","market":"sh"},{"code":"688223","name":"晶科能源","market":"sh"},
    {"code":"000066","name":"中国长城","market":"sz"},{"code":"000069","name":"华侨城A","market":"sz"},
    {"code":"000100","name":"TCL科技","market":"sz"},{"code":"000157","name":"中联重科","market":"sz"},
    {"code":"000338","name":"潍柴动力","market":"sz"},{"code":"000425","name":"徐工机械","market":"sz"},
    {"code":"000538","name":"云南白药","market":"sz"},{"code":"000625","name":"长安汽车","market":"sz"},
    {"code":"000651","name":"格力电器","market":"sz"},{"code":"000661","name":"长春高新","market":"sz"},
    {"code":"000768","name":"中航西飞","market":"sz"},{"code":"000776","name":"广发证券","market":"sz"},
    {"code":"000786","name":"北新建材","market":"sz"},{"code":"000792","name":"盐湖股份","market":"sz"},
    {"code":"000800","name":"一汽解放","market":"sz"},{"code":"000876","name":"新希望","market":"sz"},
    {"code":"000895","name":"双汇发展","market":"sz"},{"code":"000963","name":"华东医药","market":"sz"},
    {"code":"001979","name":"招商蛇口","market":"sz"},{"code":"002001","name":"新和成","market":"sz"},
    {"code":"002007","name":"华兰生物","market":"sz"},{"code":"002027","name":"分众传媒","market":"sz"},
    {"code":"002049","name":"紫光国微","market":"sz"},{"code":"002050","name":"三花智控","market":"sz"},
    {"code":"002129","name":"TCL中环","market":"sz"},{"code":"002179","name":"中航光电","market":"sz"},
    {"code":"002230","name":"科大讯飞","market":"sz"},{"code":"002236","name":"大华股份","market":"sz"},
    {"code":"002241","name":"歌尔股份","market":"sz"},{"code":"002271","name":"东方雨虹","market":"sz"},
    {"code":"002311","name":"海大集团","market":"sz"},{"code":"002371","name":"北方华创","market":"sz"},
    {"code":"002410","name":"广联达","market":"sz"},{"code":"002460","name":"赣锋锂业","market":"sz"},
    {"code":"002466","name":"天齐锂业","market":"sz"},{"code":"002493","name":"荣盛石化","market":"sz"},
    {"code":"002555","name":"三七互娱","market":"sz"},{"code":"002601","name":"龙佰集团","market":"sz"},
    {"code":"002709","name":"天赐材料","market":"sz"},{"code":"002812","name":"恩捷股份","market":"sz"},
    {"code":"002916","name":"深南电路","market":"sz"},{"code":"002920","name":"德赛西威","market":"sz"},
    {"code":"300014","name":"亿纬锂能","market":"sz"},{"code":"300122","name":"智飞生物","market":"sz"},
    {"code":"300124","name":"汇川技术","market":"sz"},{"code":"300274","name":"阳光电源","market":"sz"},
    {"code":"300308","name":"中际旭创","market":"sz"},{"code":"300316","name":"晶盛机电","market":"sz"},
    {"code":"300347","name":"泰格医药","market":"sz"},{"code":"300408","name":"三环集团","market":"sz"},
    {"code":"300413","name":"芒果超媒","market":"sz"},{"code":"300433","name":"蓝思科技","market":"sz"},
    {"code":"300450","name":"先导智能","market":"sz"},{"code":"300454","name":"深信服","market":"sz"},
    {"code":"300496","name":"中科创达","market":"sz"},{"code":"300498","name":"温氏股份","market":"sz"},
    {"code":"300628","name":"亿联网络","market":"sz"},{"code":"300760","name":"迈瑞医疗","market":"sz"},
    {"code":"300782","name":"卓胜微","market":"sz"},{"code":"300896","name":"爱美客","market":"sz"},
    {"code":"300919","name":"中伟股份","market":"sz"},{"code":"300957","name":"贝泰妮","market":"sz"},
]


def _stock_market(code):
    return "sh" if str(code).startswith(("6", "9")) else "sz"


def _normalize_stock_item(raw):
    code = str(raw.get("code") or raw.get("f12") or "").zfill(6)
    if not re.match(r"^\d{6}$", code):
        return None
    name = str(raw.get("name") or raw.get("f14") or code).strip() or code
    market = raw.get("market") or ("sh" if raw.get("f13") == 1 else _stock_market(code))
    return {
        "code": code,
        "name": name,
        "market": market,
        "industry": raw.get("industry") or raw.get("f100") or "-",
        "mv": float(raw.get("mv") or raw.get("f20") or 0),
        "pe": float(raw.get("pe") or raw.get("f9") or 0),
        "pb": float(raw.get("pb") or raw.get("f23") or 0),
    }


def _load_fundamentals():
    data = _read_json_file(os.path.join(SUPPORT_DIR, "fundamentals.json")) or {}
    if not isinstance(data, dict):
        return {}
    return data


def _fetch_eastmoney_universe():
    fields = "f12,f14,f13,f100,f20,f9,f23"
    fs = "m:0+t:6,m:0+t:80,m:1+t:2,m:1+t:23"
    url = (
        "https://push2.eastmoney.com/api/qt/clist/get"
        f"?pn=1&pz=6000&po=1&np=1&ut=bd1d9ddb04089700cf9c27f6f7426281"
        f"&fltt=2&invt=2&fid=f3&fs={fs}&fields={fields}"
    )
    try:
        r = subprocess.run(
            ["curl", "-sS", "--connect-timeout", "3", "--max-time", "10", "--noproxy", "*", url],
            capture_output=True,
            text=True,
            timeout=12,
        )
        if r.returncode != 0 or not r.stdout:
            return None
        payload = json.loads(r.stdout)
        diff = ((payload or {}).get("data") or {}).get("diff") or []
        items = [_normalize_stock_item(row) for row in diff]
        items = [item for item in items if item]
        if not items:
            return None
        snapshot = {
            "items": items,
            "total": len(items),
            "source": "eastmoney_clist",
            "updated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }
        _write_json_file(UNIVERSE_CACHE_PATH, snapshot)
        return snapshot
    except Exception:
        return None


def _fetch_akshare_universe():
    try:
        import contextlib
        import akshare as ak

        with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
            df = ak.stock_info_a_code_name()
        fundamentals = _load_fundamentals()
        core_names = {s["code"]: s["name"] for s in STOCK_POOL}
        items = []
        seen = set()
        for row in df.to_dict("records"):
            code = str(row.get("code") or "").zfill(6)
            info = fundamentals.get(code, {}) or {}
            item = _normalize_stock_item({
                "code": code,
                "name": core_names.get(code) or row.get("name") or code,
                "market": _stock_market(code),
                **info,
            })
            if item and item["code"] not in seen:
                seen.add(item["code"])
                items.append(item)
        if not items:
            return None
        snapshot = {
            "items": items,
            "total": len(items),
            "source": "akshare_stock_info_a_code_name",
            "updated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }
        _write_json_file(UNIVERSE_CACHE_PATH, snapshot)
        return snapshot
    except Exception:
        return None


def load_all_stock_pool(refresh=False):
    if refresh:
        fresh = _fetch_eastmoney_universe() or _fetch_akshare_universe()
        if fresh:
            return fresh
    cached = _read_json_file(UNIVERSE_CACHE_PATH)
    if cached and isinstance(cached.get("items"), list) and cached.get("items"):
        return cached

    fresh = _fetch_akshare_universe()
    if fresh:
        return fresh

    fundamentals = _load_fundamentals()
    core_names = {s["code"]: s["name"] for s in STOCK_POOL}
    items = []
    seen = set()
    for code, info in fundamentals.items():
        item = _normalize_stock_item({
            "code": code,
            "name": core_names.get(code, code),
            "market": _stock_market(code),
            **(info or {}),
        })
        if item and item["code"] not in seen:
            seen.add(item["code"])
            items.append(item)
    for stock in STOCK_POOL:
        item = _normalize_stock_item(stock)
        if item and item["code"] not in seen:
            seen.add(item["code"])
            items.append(item)
    items.sort(key=lambda x: (x.get("mv") or 0), reverse=True)
    return {
        "items": items,
        "total": len(items),
        "source": "fundamentals_plus_core_fallback",
        "updated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }


def resolve_stock_universe(scope="core", q="", limit=0, offset=0, refresh=False):
    if scope == "all":
        snapshot = load_all_stock_pool(refresh=refresh)
        items = snapshot.get("items", [])
        source = snapshot.get("source", "unknown")
        updated = snapshot.get("updated")
    else:
        items = STOCK_POOL
        source = "core_static_pool"
        updated = None

    q = (q or "").strip().lower()
    if q:
        items = [
            item for item in items
            if q in item.get("code", "").lower() or q in item.get("name", "").lower()
        ]
    total = len(items)
    offset = max(0, int(offset or 0))
    limit = max(0, int(limit or 0))
    page_items = items[offset: offset + limit] if limit else items[offset:]
    return {
        "items": page_items,
        "total": total,
        "offset": offset,
        "limit": limit,
        "scope": scope,
        "source": source,
        "updated": updated,
    }


def stock_name_map(scope="core"):
    return {s["code"]: s["name"] for s in resolve_stock_universe(scope).get("items", [])}


def stock_name(code):
    for stock in STOCK_POOL:
        if stock["code"] == code:
            return stock["name"]
    for stock in load_all_stock_pool().get("items", []):
        if stock["code"] == code:
            return stock.get("name") or code
    return code


def _clamp(value, low, high):
    return max(low, min(high, value))


def _manual_decision_context():
    data = _read_json_file(DECISION_CONTEXT_PATH) or {}
    if not isinstance(data, dict):
        data = {}
    defaults = {
        "monetary_policy": {"label": "货币政策", "score": 0.0, "source": "not_configured", "message": "未配置实时政策数据，按中性处理"},
        "currency_value": {"label": "货币价值", "score": 0.0, "source": "not_configured", "message": "未配置人民币/通胀/实际利率数据，按中性处理"},
        "fiscal_policy": {"label": "财政政策", "score": 0.0, "source": "not_configured", "message": "未配置财政支出/专项债等数据，按中性处理"},
        "main_fund_flow": {"label": "主力资金", "score": 0.0, "source": "not_configured", "message": "未配置主力净流入，按中性处理"},
        "retail_fund_flow": {"label": "散户资金", "score": 0.0, "source": "not_configured", "message": "未配置散户交易拥挤度，按中性处理"},
    }
    for key, value in data.items():
        if key in defaults and isinstance(value, dict):
            defaults[key].update(value)
    return defaults


def build_decision_context(prediction):
    feat = (prediction or {}).get("features_used", {}) or {}
    manual = _manual_decision_context()
    dynamic = {
        "market_regime": {
            "label": "市场状态",
            "score": _clamp(float(feat.get("market_regime") or 0), -1, 1),
            "source": "hs300_ma20_ret20",
            "message": "沪深300 MA20 与 20日收益率",
        },
        "market_breadth": {
            "label": "赚钱效应",
            "score": _clamp((float(feat.get("market_breadth") or 0.5) - 0.5) * 2, -1, 1),
            "source": "cross_section_cache",
            "message": "全市场/股票池上涨家数比例",
        },
        "northbound_flow": {
            "label": "北向资金",
            "score": _clamp(float(feat.get("nb_net_flow") or 0) / 100.0, -1, 1),
            "source": "northbound_flow_cache",
            "message": "北向资金净流入强弱",
        },
        "margin_flow": {
            "label": "杠杆资金",
            "score": _clamp(float(feat.get("margin_buy_ratio") or 0), -1, 1),
            "source": "margin_cache",
            "message": "融资买入/余额变化",
        },
    }
    factors = {**manual, **dynamic}
    scores = [float(v.get("score") or 0) for v in factors.values()]
    aggregate = sum(scores) / len(scores) if scores else 0.0
    return {
        "aggregate_score": round(_clamp(aggregate, -1, 1), 3),
        "risk_multiplier": round(_clamp(1 + aggregate * 0.2, 0.65, 1.25), 3),
        "factors": factors,
        "config_path": DECISION_CONTEXT_PATH,
    }


def apply_decision_overlay(prediction):
    if not isinstance(prediction, dict) or prediction.get("error"):
        return prediction
    context = build_decision_context(prediction)
    prediction["decision_context"] = context
    pos = prediction.get("suggested_position")
    if isinstance(pos, dict):
        reasons = list(pos.get("reasons") or [])
        weight = float(pos.get("weight") or 0)
        if weight > 0 and context["risk_multiplier"] < 0.9:
            weight = round(weight * context["risk_multiplier"], 2)
            reasons.append("货币/财政/资金/市场环境偏弱，压低风险预算")
        elif weight > 0 and context["risk_multiplier"] > 1.1:
            reasons.append("宏观与资金环境支持，保留模型仓位")
        pos["weight"] = weight
        pos["reasons"] = reasons
        prediction["suggested_position"] = pos
    return prediction


def apply_decision_overlay_to_payload(data):
    if not isinstance(data, dict):
        return data
    predictions = data.get("predictions")
    if isinstance(predictions, list):
        data = dict(data)
        data["predictions"] = [apply_decision_overlay(item) for item in predictions]
    return data


def _future_trading_dates(last_date, horizon):
    try:
        cursor = datetime.strptime(last_date, "%Y-%m-%d")
    except Exception:
        cursor = datetime.now()
    dates = []
    while len(dates) < horizon:
        cursor += timedelta(days=1)
        if cursor.weekday() < 5:
            dates.append(cursor.strftime("%Y-%m-%d"))
    return dates


def build_forecast_path(kl, prediction, horizon=5):
    horizon = max(1, min(20, int(horizon or 5)))
    last_date, _, last_close, _, _, _ = kl[-1]
    feat = (prediction or {}).get("features_used", {}) or {}
    class_probs = (prediction or {}).get("class_probabilities", {}) or {}
    out_prob = float((prediction or {}).get("outperform_prob") or 0.5)
    excess_map = {"强跑输": -0.04, "小跑输": -0.012, "小跑赢": 0.012, "强跑赢": 0.04}
    class_excess = sum(float(class_probs.get(label, 0)) * ret for label, ret in excess_map.items())
    edge_excess = (out_prob - 0.5) * 0.08
    expected_excess_5d = class_excess * 0.6 + edge_excess * 0.4
    benchmark_5d = _clamp(float(feat.get("bench_ret_20") or 0) / 100.0 * 0.25, -0.035, 0.035)
    context = build_decision_context(prediction)
    macro_adjustment = context["aggregate_score"] * 0.008
    expected_5d = _clamp(benchmark_5d + expected_excess_5d + macro_adjustment, -0.12, 0.12)
    daily_vol = _clamp(float(feat.get("hist_vol_60") or feat.get("volatility_20") or 0.02), 0.006, 0.08)
    atr = _clamp(float(feat.get("atr_14_pct") or 2.0) / 100.0, 0.004, 0.12)
    band_5d = 1.28 * daily_vol * math.sqrt(horizon) * context["risk_multiplier"]
    dates = _future_trading_dates(last_date, horizon)
    recent_actual = [
        {
            "date": item[0],
            "open": round(float(item[1]), 2),
            "close": round(float(item[2]), 2),
            "low": round(float(item[4]), 2),
            "high": round(float(item[3]), 2),
        }
        for item in kl[-FORECAST_HISTORY_DAYS:]
    ]
    candles = []
    base_line = []
    bull_line = []
    bear_line = []
    prev_close = float(last_close)
    for idx, date in enumerate(dates, start=1):
        step = idx / horizon
        close = float(last_close) * (1 + expected_5d * step)
        open_price = prev_close
        intraday = max(atr * 0.45, daily_vol * 0.75, 0.006)
        high = max(open_price, close) * (1 + intraday / 2)
        low = min(open_price, close) * (1 - intraday / 2)
        candles.append({
            "date": date,
            "open": round(open_price, 2),
            "close": round(close, 2),
            "low": round(low, 2),
            "high": round(high, 2),
        })
        base_line.append(round(close, 2))
        bull_line.append(round(float(last_close) * (1 + (expected_5d + band_5d) * step), 2))
        bear_line.append(round(float(last_close) * (1 + (expected_5d - band_5d) * step), 2))
        prev_close = close
    return {
        "last_date": last_date,
        "last_close": round(float(last_close), 2),
        "history_window_target_days": FORECAST_HISTORY_DAYS,
        "history_window_days": min(len(kl), FORECAST_HISTORY_DAYS),
        "horizon_days": horizon,
        "expected_5d_return_pct": round(expected_5d * 100, 2),
        "expected_excess_5d_pct": round(expected_excess_5d * 100, 2),
        "benchmark_5d_assumption_pct": round(benchmark_5d * 100, 2),
        "uncertainty_band_pct": round(band_5d * 100, 2),
        "recent_actual": recent_actual,
        "candles": candles,
        "base_line": base_line,
        "bull_line": bull_line,
        "bear_line": bear_line,
        "decision_context": context,
        "disclaimer": "预测K线是概率情景路径，不是确定未来价格。",
    }

CACHE_WARMUP_STARTED = False


def warm_kline_cache():
    global CACHE_WARMUP_STARTED
    if CACHE_WARMUP_STARTED:
        return
    CACHE_WARMUP_STARTED = True
    codes = [s["code"] for s in STOCK_POOL]

    def warm_one(code):
        if read_kline_cache(code, 500, allow_stale=True):
            return
        fetch_kline(code, 500)

    with ThreadPoolExecutor(max_workers=16) as pool:
        for _ in as_completed([pool.submit(warm_one, code) for code in codes]):
            pass


@app.on_event("startup")
def start_cache_warmup():
    threading.Thread(target=warm_kline_cache, daemon=True).start()


PRESET_STRATEGIES = {
    "rsi":{"name":"RSI反转","desc":"RSI超卖买入，超买卖出","params":{"oversold":25,"overbought":80},
        "param_meta":{"oversold":{"type":"range","min":15,"max":35,"desc":"超卖阈值"},"overbought":{"type":"range","min":65,"max":85,"desc":"超买阈值"}}},
    "ma_cross":{"name":"均线金叉","desc":"快线上穿慢线买，下穿卖","params":{"fast":5,"slow":20},
        "param_meta":{"fast":{"type":"range","min":3,"max":15,"desc":"快线周期"},"slow":{"type":"range","min":15,"max":90,"desc":"慢线周期"}}},
    "turtle":{"name":"海龟突破","desc":"突破N日高买，跌破M日低卖","params":{"entry":20,"exit_":10},
        "param_meta":{"entry":{"type":"range","min":10,"max":60,"desc":"突破周期"},"exit_":{"type":"range","min":5,"max":30,"desc":"退出周期"}}},
    "momentum":{"name":"动量追涨","desc":"强势追涨，弱势止损","params":{"ma_period":60,"mom_threshold":0.10},
        "param_meta":{"ma_period":{"type":"range","min":30,"max":120,"desc":"均线周期"},"mom_threshold":{"type":"range","min":0.05,"max":0.25,"step":0.01,"desc":"动量阈值"}}},
    "rsi_trailing":{"name":"RSI+移动止盈","desc":"RSI超卖买，移动止盈卖","params":{"oversold":25,"trail_pct":0.08,"rsi_exit":75},
        "param_meta":{"oversold":{"type":"range","min":15,"max":35,"desc":"超卖阈值"},"trail_pct":{"type":"range","min":0.03,"max":0.15,"step":0.01,"desc":"移动止盈%"},"rsi_exit":{"type":"range","min":60,"max":85,"desc":"RSI退出"}}},
    "consensus":{"name":"三票共识","desc":"RSI+MA交叉+动量投票，≥2票买入","params":{"oversold":25,"fast":5,"slow":20,"ma_period":60,"mom_threshold":0.10,"min_votes":2,"overbought":80},
        "param_meta":{"oversold":{"type":"range","min":15,"max":35,"desc":"RSI超卖阈值"},"min_votes":{"type":"range","min":1,"max":3,"desc":"最少同意票数"}}},
    "adaptive":{"name":"自适应(ADX判市)","desc":"ADX>25追趋势，ADX<15做均值回复，中间等金叉。全自动切换+ATR止损。","params":{"adx_trend":25,"adx_choppy":15,"rsi_os":22,"rsi_ob":78},
        "param_meta":{"adx_trend":{"type":"range","min":20,"max":35,"desc":"ADX趋势阈值"},"adx_choppy":{"type":"range","min":10,"max":20,"desc":"ADX震荡阈值"},"rsi_os":{"type":"range","min":15,"max":30,"desc":"RSI超卖"},"rsi_ob":{"type":"range","min":70,"max":88,"desc":"RSI超买"}}},
}

# ═══════════ WebSocket 连接池 ═══════════
ws_clients = set()
ALERT_COOLDOWN = {}
ALERT_COOLDOWN_SECONDS = 20 * 60


def build_realtime_alert(code, info):
    price = info.get("price") or 0
    yest_close = info.get("yest_close") or 0
    change = (price - yest_close) / yest_close * 100 if yest_close > 0 else 0
    kl = fetch_kline(code, 120)
    if not kl:
        return None

    closes = [k[2] for k in kl]
    highs = [k[3] for k in kl]
    lows = [k[4] for k in kl]
    closes_live = closes + [price]
    rsi = calc_rsi(closes_live)
    ma20 = calc_ma(closes_live, 20)
    ma60 = calc_ma(closes_live, 60)
    ma20_dev = (price - ma20) / ma20 * 100 if ma20 > 0 else 0
    ma60_dev = (price - ma60) / ma60 * 100 if ma60 > 0 else 0
    trend20 = _trend_angle_from_closes(closes_live, 20)
    atr = _atr_pct_from_kline(kl, 14)
    h20 = max(highs[-20:]) if len(highs) >= 20 else max(highs)
    l20 = min(lows[-20:]) if len(lows) >= 20 else min(lows)

    signal = None
    severity = "low"
    action = "观察"
    reason = ""

    if change <= -4 or (ma60_dev <= -6 and trend20 < -8):
        signal = "RISK_OFF"
        severity = "high"
        action = "减仓/不买"
        reason = "日内跌幅或MA60趋势破位"
    elif rsi >= 82 and change > 1:
        signal = "TAKE_PROFIT"
        severity = "medium"
        action = "止盈/别追"
        reason = "RSI极热且上涨，容易冲高回落"
    elif price >= h20 * 0.995 and rsi < 78 and trend20 > 5 and ma20_dev > 0:
        signal = "BREAKOUT_WATCH"
        severity = "medium"
        action = "突破观察"
        reason = "接近20日新高且趋势向上"
    elif rsi <= 24 and ma60_dev > -8 and change > -3:
        signal = "REVERSAL_WATCH"
        severity = "medium"
        action = "小仓观察"
        reason = "RSI超卖但未深度破位"
    elif rsi <= 20 and (ma60_dev <= -8 or trend20 < -8):
        signal = "FALLING_KNIFE"
        severity = "high"
        action = "不接飞刀"
        reason = "超卖叠加趋势破位"
    elif atr >= 5 and abs(change) >= 2:
        signal = "HIGH_VOL"
        severity = "medium"
        action = "降低仓位"
        reason = "ATR和日内波动同步放大"

    if not signal:
        return None

    key = f"{code}:{signal}"
    now = time.time()
    if now - ALERT_COOLDOWN.get(key, 0) < ALERT_COOLDOWN_SECONDS:
        return None
    ALERT_COOLDOWN[key] = now

    return {
        "code": code,
        "name": info.get("name"),
        "price": price,
        "change": round(change, 2),
        "rsi": round(rsi, 1),
        "ma20_dev": round(ma20_dev, 2),
        "ma60_dev": round(ma60_dev, 2),
        "trend20": trend20,
        "atr": atr,
        "signal": signal,
        "severity": severity,
        "action": action,
        "reason": reason,
    }

@app.websocket("/ws/monitor")
async def ws_monitor(ws: WebSocket):
    await ws.accept()
    ws_clients.add(ws)
    try:
        while True:
            await asyncio.sleep(5)
            try:
                codes = ['sh600519','sh600036','sz000001','sz300750','sh601318','sh600276','sz002415','sh603259','sz000333','sz000858','sh601899','sz300059']
                rt = fetch_realtime(codes)
                alerts = []
                for raw_code, info in rt.items():
                    code = raw_code[2:]
                    alert = build_realtime_alert(code, info)
                    if alert:
                        alerts.append(alert)
                await ws.send_json({'time':datetime.now().strftime('%H:%M:%S'),'data':alerts})
            except Exception:
                break
    except WebSocketDisconnect:
        pass
    finally:
        ws_clients.discard(ws)

# ═══════════ API ═══════════

@app.get("/api/stocks")
def api_stocks():
    return STOCK_POOL


@app.get("/api/stocks/universe")
def api_stocks_universe(
    q: str = Query(""),
    limit: int = Query(100),
    offset: int = Query(0),
    refresh: bool = Query(False),
):
    return resolve_stock_universe("all", q=q, limit=limit, offset=offset, refresh=refresh)


@app.post("/api/stocks/universe/refresh")
def api_stocks_universe_refresh():
    snapshot = load_all_stock_pool(refresh=True)
    return {
        "total": snapshot.get("total", 0),
        "source": snapshot.get("source", "unknown"),
        "updated": snapshot.get("updated"),
    }

@app.get("/api/stock/{code}/history")
def api_history(code: str, days: int = HISTORY_WINDOW_DAYS):
    days = max(1, min(int(days or HISTORY_WINDOW_DAYS), HISTORY_WINDOW_DAYS))
    kl = fetch_kline(code, days)
    if not kl: return JSONResponse({"error":"无数据"},404)
    closes=[k[2] for k in kl]
    return {"code":code,"total_days":len(kl),
        "history_window_days": HISTORY_WINDOW_DAYS,
        "recent":[{"date":k[0],"open":k[1],"close":k[2],"high":k[3],"low":k[4],"volume":k[5]} for k in kl[-HISTORY_WINDOW_DAYS:]],
        "last_close":closes[-1],"rsi":round(calc_rsi(closes),1),
        "ma20":round(calc_ma(closes,20),2),"ma60":round(calc_ma(closes,60),2),
        "change_20d":round((closes[-1]-closes[-20])/closes[-20]*100,2) if len(closes)>=20 else 0}


@app.get("/api/stock/{code}/detail")
def api_stock_detail(code: str):
    """单只股票完整详情"""
    kl = fetch_kline(code, HISTORY_WINDOW_DAYS)
    if not kl: return JSONResponse({"error":"无数据"}, 404)
    closes = [k[2] for k in kl]
    r = calc_rsi(closes)
    name = stock_name(code)

    # 计算多个指标
    def trend_angle(prices, n):
        if len(prices) < n: return 0
        y = prices[-n:]; x = list(range(n))
        n_ = n; sxy = sum(x[i]*y[i] for i in range(n_))
        sx = sum(x); sy = sum(y); sxx = sum(x[i]*x[i] for i in range(n_))
        slope = (n_*sxy - sx*sy) / (n_*sxx - sx*sx + 0.0001)
        return math.degrees(math.atan(slope))

    returns = [(closes[i]-closes[i-1])/closes[i-1] for i in range(1, len(closes))]
    vol = (sum(rr*rr for rr in returns)/len(returns))**0.5 * 100 if returns else 0

    # 回溯RSI序列 — 预计算涨跌幅避免 O(n²)
    rsi_series = [50]*14
    gains = [max(closes[i]-closes[i-1],0) for i in range(1, len(closes))]
    losses = [max(closes[i-1]-closes[i],0) for i in range(1, len(closes))]
    for i in range(len(gains)):
        w = min(14, i+1)
        ag = sum(gains[max(0,i-13):i+1])/w
        al = sum(losses[max(0,i-13):i+1])/w
        rsi_series.append(100-100/(1+ag/(al+0.0001)))

    recent = [{"date": kl[i][0], "open": kl[i][1], "close": kl[i][2], "high": kl[i][3], "low": kl[i][4], "volume": kl[i][5], "rsi": round(rsi_series[i],1)} for i in range(max(0,len(kl)-HISTORY_WINDOW_DAYS), len(kl))]

    ma20 = sum(closes[-20:])/20 if len(closes)>=20 else closes[-1]
    ma60 = sum(closes[-60:])/60 if len(closes)>=60 else closes[-1]

    return {
        "code": code, "name": name,
        "history_window_days": HISTORY_WINDOW_DAYS,
        "total_days": len(kl),
        "last_close": closes[-1], "rsi": round(r,1),
        "trend_60": round(trend_angle(closes, 60), 1),
        "trend_120": round(trend_angle(closes, 120), 1),
        "trend_252": round(trend_angle(closes, min(252, len(closes))), 1),
        "volatility": round(vol, 2),
        "ma20": round(ma20, 2), "ma60": round(ma60, 2),
        "high_60d": round(max(closes[-60:]), 2), "low_60d": round(min(closes[-60:]), 2),
        "recent": recent
    }

@app.get("/api/stock/{code}/realtime")
def api_realtime(code: str):
    rt_key = f'sh{code}' if code.startswith(('6','9')) else f'sz{code}'
    rt = fetch_realtime([rt_key])
    info = rt.get(rt_key)
    if not info: return JSONResponse({"error":"获取失败"},500)
    kl = fetch_kline(code, 100)
    if kl:
        closes=[k[2] for k in kl]; closes.append(info['price'])
        r=calc_rsi(closes)
    else: r=50
    return {**info,"rsi":round(r,1),"change_pct":round((info['price']-info['yest_close'])/info['yest_close']*100,2)}

@app.get("/api/dashboard")
def api_dashboard():
    codes = [s["code"] for s in STOCK_POOL[:30]]
    items = []; oversold = 0; overbought = 0

    def analyze_one(code):
        kl = fetch_kline(code, 100)
        if not kl: return None
        closes = [k[2] for k in kl]
        name = next((s["name"] for s in STOCK_POOL if s["code"]==code), code)
        r = round(calc_rsi(closes), 1)
        chg = round((closes[-1]-closes[-20])/closes[-20]*100, 2) if len(closes)>=20 else 0
        return {"code":code,"name":name,"close":closes[-1],"rsi":r,"change_20d":chg,"ma20":round(calc_ma(closes,20),2),"signal":"BUY" if r<25 else "SELL" if r>80 else None}

    with ThreadPoolExecutor(max_workers=min(16, max(1, len(codes)))) as pool:
        for future in as_completed([pool.submit(analyze_one, code) for code in codes]):
            item = future.result()
            if item:
                items.append(item)
                if item["rsi"] < 25: oversold += 1
                if item["rsi"] > 80: overbought += 1

    items.sort(key=lambda x: x["rsi"])
    return {"items":items,"oversold":oversold,"overbought":overbought,"sentiment":"偏弱·机会" if oversold>0 else "偏强·谨慎" if overbought>0 else "中性","total":len(items)}

@app.get("/api/strategies")
def api_strategies():
    result = [{"id":k,"name":v["name"],"desc":v["desc"],"params":v["params"],"param_meta":v.get("param_meta",{})} for k,v in PRESET_STRATEGIES.items()]
    result.append({"id":"custom","name":"自定义策略","desc":"编写你自己的买卖条件","params":{"buy_condition":"rsi_14 < 25","sell_condition":"rsi_14 > 80"},"param_meta":{"buy_condition":{"type":"code","desc":"买入条件(Python表达式)"},"sell_condition":{"type":"code","desc":"卖出条件(Python表达式)"}}})
    return result

@app.post("/api/backtest")
def api_backtest(data: dict):
    codes = list(dict.fromkeys(data.get("codes",[])))
    sid = data.get("strategy","rsi")
    params = data.get("params",{})
    if sid=="custom" and "buy_condition" not in params:
        return JSONResponse({"error":"自定义策略需要 buy_condition 和 sell_condition"},400)
    if sid in PRESET_STRATEGIES:
        params = {**PRESET_STRATEGIES[sid]["params"], **params}
    else:
        params["strategy"] = "custom"

    # ── Phase 2: 选股质量预过滤 ──
    enable_filter = data.get("filter_stocks", True)  # 默认开启
    filtered_out = []
    if enable_filter and len(codes) >= 5:
        quality_codes = []
        for code in codes:
            kl = fetch_kline(code, 120)
            if not kl or len(kl) < 60:
                continue
            closes = [k[2] for k in kl]
            highs = [k[3] for k in kl]; lows = [k[4] for k in kl]
            # 趋势过滤器：排除明显下跌趋势
            trend_60 = _trend_angle_from_closes(closes, 60)
            if trend_60 < -8:
                filtered_out.append(f"{code}(趋势{trend_60:.0f}°)")
                continue
            # 价格过滤器：排除高价股（>500元，资金过度集中）
            if closes[-1] > 500:
                filtered_out.append(f"{code}(高价{closes[-1]:.0f})")
                continue
            # 波动率过滤器：排除极端波动（ATR>8%）
            atr_pct = _atr_pct_from_kline(kl, 14)
            if atr_pct > 8:
                filtered_out.append(f"{code}(高波{atr_pct:.1f}%)")
                continue
            quality_codes.append(code)
        if quality_codes:
            codes = quality_codes

    results=[]
    total_pnl=0
    max_workers = min(32, max(1, len(codes)))
    with ThreadPoolExecutor(max_workers=max_workers) as pool:
        future_map = {pool.submit(run_backtest_detailed, code, {**params, "strategy":sid}): code for code in codes}
        for future in as_completed(future_map):
            code = future_map[future]
            try:
                r = future.result()
            except Exception:
                r = None
            if r:
                results.append({"code":code,**r})
                total_pnl += r["pnl"]
    results.sort(key=lambda x:x["pnl"], reverse=True)
    return {"strategy":PRESET_STRATEGIES.get(sid,{}).get("name","自定义"),"params":params,"results":results,
        "total_pnl":total_pnl,"total_trades":sum(r["trades"] for r in results),
        "avg_win_rate":round(sum(r["win_rate"] for r in results)/len(results),1) if results else 0,
        "avg_sharpe":round(sum(r["sharpe"] for r in results)/len(results),2) if results else 0,
        "max_dd":round(max(r["max_drawdown"] for r in results),2) if results else 0,
        "filtered_out": filtered_out}

@app.post("/api/backtest/export")
def api_backtest_export(data: dict):
    """导出回测结果为 CSV"""
    codes = data.get("codes",[]); sid = data.get("strategy","rsi"); params = data.get("params",{})
    if sid in PRESET_STRATEGIES: params = {**PRESET_STRATEGIES[sid]["params"], **params}
    else: params["strategy"]="custom"
    output = io.StringIO()
    # UTF-8 BOM for Excel compatibility with Chinese characters
    output.write('﻿')
    w = csv.writer(output)
    w.writerow(["股票","买入日","买入价","卖出日","卖出价","盈亏","收益率%","持仓天数","最大浮盈%"])
    for code in codes:
        r = run_backtest_detailed(code, {**params, "strategy":sid})
        if r:
            for t in r["trade_list"]:
                w.writerow([code,t['buy_date'],t['buy_price'],t['sell_date'],t['sell_price'],t['pnl'],t['pct'],t['holding_days'],t['peak_dd']])
    output.seek(0)
    csv_bytes = output.getvalue().encode('utf-8-sig')
    return Response(content=csv_bytes, media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename=backtest_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"})

@app.post("/api/optimize")
def api_optimize(data: dict):
    """自动寻优：对所有策略做参数网格搜索，返回排名（精简版，每个参数只测min/max)"""
    codes = data.get("codes", [s["code"] for s in STOCK_POOL[:10]])
    if not codes:
        return JSONResponse({"error": "请选择至少1只股票"}, 400)

    strategies_to_test = data.get("strategies") or list(PRESET_STRATEGIES.keys())
    strategies_to_test = [s for s in strategies_to_test if s in PRESET_STRATEGIES]
    if not strategies_to_test:
        return JSONResponse({"error": "无有效策略"}, 400)

    started_at = time.time()
    all_results = []
    total_combos = 0
    errors = []

    for sid in strategies_to_test:
        strat = PRESET_STRATEGIES[sid]
        meta = strat.get("param_meta", {})
        params = dict(strat["params"])

        combos = _generate_param_combos(params, meta)
        total_combos += len(combos)

        for combo in combos:
            merged = {**params, **combo, "strategy": sid}
            total_pnl = 0; total_trades = 0; total_wins = 0
            max_dd_all = 0; sharpes = []; win_rates = []
            code_results = []

            try:
                with ThreadPoolExecutor(max_workers=min(8, max(1, len(codes)))) as pool:
                    future_map = {pool.submit(run_backtest_detailed, code, merged): code for code in codes}
                    for future in as_completed(future_map):
                        code = future_map[future]
                        try:
                            r = future.result()
                        except Exception as e:
                            errors.append(f"{sid}/{code}: {e}")
                            r = None
                        if r and r.get("trades", 0) > 0:
                            total_pnl += r["pnl"]
                            total_trades += r["trades"]
                            total_wins += r["wins"]
                            max_dd_all = max(max_dd_all, r["max_drawdown"])
                            sharpes.append(r["sharpe"])
                            win_rates.append(r["win_rate"])
                            code_results.append({"code": code, "pnl": r["pnl"], "trades": r["trades"],
                                                "win_rate": r["win_rate"], "sharpe": r["sharpe"]})
            except Exception as e:
                errors.append(f"{sid}/{combo}: {e}")
                continue

            if total_trades >= 3:
                avg_sharpe = sum(sharpes) / len(sharpes) if sharpes else 0
                avg_win_rate = sum(win_rates) / len(win_rates) if win_rates else 0
                composite = avg_sharpe * 0.5 + avg_win_rate / 100 * 0.2 + min(total_trades / 20, 1) * 0.15
                if max_dd_all > 0:
                    composite += (1 - min(max_dd_all / 30, 1)) * 0.15

                all_results.append({
                    "strategy": sid,
                    "strategy_name": strat["name"],
                    "params": combo,
                    "total_pnl": total_pnl,
                    "total_trades": total_trades,
                    "total_wins": total_wins,
                    "avg_win_rate": round(avg_win_rate, 1),
                    "avg_sharpe": round(avg_sharpe, 2),
                    "max_drawdown": round(max_dd_all, 2),
                    "composite_score": round(composite, 3),
                    "stocks_tested": len(code_results),
                    "top_codes": sorted(code_results, key=lambda x: x["pnl"], reverse=True)[:5],
                })

    all_results.sort(key=lambda x: x["composite_score"], reverse=True)
    return {
        "results": all_results[:20],
        "total_combinations": total_combos,
        "strategies_tested": len(strategies_to_test),
        "stocks_tested": len(codes),
        "elapsed": round(time.time() - started_at, 1),
        "best": all_results[0] if all_results else None,
        "errors": errors[:10] if errors else [],
    }


def _generate_param_combos(params, meta):
    """生成参数网格：每个参数只取 min/max 两端（速度优先）"""
    axes = {}
    for key, config in meta.items():
        v = params.get(key)
        mn = config.get("min"); mx = config.get("max")
        step = config.get("step")
        if mn is not None and mx is not None:
            if step and step < (mx - mn) * 0.5:
                # 有步长且步长合理：取 min, mid, max
                mid = round((mn + mx) / 2, 4)
                axes[key] = [mn, mid, mx]
            else:
                # 无步长或步长太大：只取 min/max 两端
                axes[key] = [mn, mx]
        else:
            axes[key] = [v]

    keys = list(axes.keys())
    if not keys:
        return [{}]
    combos = [{}]
    for k in keys:
        new_combos = []
        for combo in combos:
            for val in axes[k]:
                c = dict(combo)
                # 保持原始类型
                orig = params.get(k)
                if isinstance(orig, int) and not isinstance(val, int):
                    c[k] = int(val)
                elif isinstance(orig, float) and not isinstance(val, float):
                    c[k] = float(val)
                else:
                    c[k] = val
                new_combos.append(c)
        combos = new_combos
    return combos

@app.post("/api/screen")
def api_screen(data: dict):
    started_at = time.time()
    if data.get("codes"):
        universe = {"items": [], "source": "request_codes", "total": len(data.get("codes") or [])}
        codes = data.get("codes") or []
    else:
        scope = data.get("universe_scope", "all")
        max_universe = int(data.get("max_universe", 500) or 500)
        universe = resolve_stock_universe(scope, limit=max_universe)
        codes = [s["code"] for s in universe.get("items", [])]
    name_map = stock_name_map("all")

    def analyze_one(code):
        kl = fetch_kline(code,300)
        if not kl or len(kl)<60: return None
        closes=[k[2] for k in kl]
        def trend_angle(n):
            y=closes[-n:]; x=list(range(n))
            n_=n; sxy=sum(xi*yi for xi,yi in zip(x,y)); sx,sy=sum(x),sum(y); sxx=sum(xi*xi for xi in x)
            slope=(n_*sxy-sx*sy)/(n_*sxx-sx*sx+0.0001)
            return math.degrees(math.atan(slope))
        a60=trend_angle(60); a120=trend_angle(120)
        returns=[(closes[i]-closes[i-1])/closes[i-1] for i in range(1,len(closes))]
        vol=(sum(r*r for r in returns)/len(returns))**0.5*100
        r=calc_rsi(closes)
        score=a60*0.35+a120*0.25+max(0,50-abs(r-50))*0.2+max(0,50-abs(vol-2)*15)*0.2
        return {"code":code,"name":name_map.get(code,code),"score":round(score,1),"trend_60":round(a60,1),"trend_120":round(a120,1),"volatility":round(vol,2),"rsi":round(r,1),"last_close":closes[-1]}

    results=[]
    codes = list(dict.fromkeys(codes))
    max_workers = min(48, max(1, len(codes)))
    with ThreadPoolExecutor(max_workers=max_workers) as pool:
        for future in as_completed([pool.submit(analyze_one, code) for code in codes]):
            try:
                item = future.result()
            except Exception:
                item = None
            if item: results.append(item)
    results.sort(key=lambda x:x["score"],reverse=True)
    return {
        "results": results,
        "total": len(codes),
        "universe_total": universe.get("total", len(codes)),
        "universe_source": universe.get("source", "unknown"),
        "analyzed": len(results),
        "skipped": max(0, len(codes) - len(results)),
        "elapsed": round(time.time() - started_at, 2)
    }

@app.post("/api/analyze")
def api_analyze(data: dict):
    api_key = load_anthropic_api_key()
    if not api_key: return JSONResponse({"error":"需要 ANTHROPIC_API_KEY"},400)
    text = data.get("orders_text","")
    if not text: return JSONResponse({"error":"请提供回测数据"},400)
    import requests as req
    payload = {"model":"claude-sonnet-4-6","max_tokens":2048,"temperature":0.3,"messages":[{"role":"user","content":f"分析以下回测结果，输出JSON(只输出JSON): {{{{\"summary\":\"总体评价\",\"loss_patterns\":\"亏损共性\",\"win_patterns\":\"盈利共性\",\"suggestions\":\"优化建议\",\"next_steps\":\"下一步\"}}}}\n\n{text}"}]}
    try:
        r=req.post("https://api.anthropic.com/v1/messages",headers={"x-api-key":api_key,"anthropic-version":"2023-06-01","content-type":"application/json"},json=payload,timeout=45)
        r.raise_for_status()
        resp=r.json()["content"][0]["text"]
        m=re.search(r'\{.*\}',resp,re.DOTALL)
        return json.loads(m.group()) if m else {"raw":resp}
    except Exception as e:
        return JSONResponse({"error":str(e)},500)

@app.post("/api/report")
def api_report(data: dict):
    codes=data.get("codes",["600519","000001","300750"])
    rt_codes=[f'sh{c}' if c.startswith(('6','9')) else f'sz{c}' for c in codes]
    rt=fetch_realtime(rt_codes)
    lines=[f"# A股市场观察 {datetime.now().strftime('%Y-%m-%d')}\n"]
    for code in codes:
        rt_key=f'sh{code}' if code.startswith(('6','9')) else f'sz{code}'
        info=rt.get(rt_key)
        if not info: continue
        kl=fetch_kline(code,200)
        rsi_v=50; ma20=info['price']
        if kl:
            closes=[k[2] for k in kl]; closes.append(info['price'])
            rsi_v=round(calc_rsi(closes),1); ma20=calc_ma(closes,20)
        chg=(info['price']-info['yest_close'])/info['yest_close']*100
        sig="🔵超卖" if rsi_v<25 else "🔴超买" if rsi_v>80 else "🟡偏高" if rsi_v>70 else "⚪正常"
        lines.append(f"## {info['name']} ({code})")
        lines.append(f"- 现价:{info['price']:.2f} | 涨跌:{chg:+.2f}%")
        lines.append(f"- RSI(14):{rsi_v} {sig} | MA20:{ma20:.2f}")
        lines.append("")
    report="\n".join(lines)
    try:
        from docx import Document
        doc=Document(); doc.add_heading(f'A股观察 {datetime.now().strftime("%Y-%m-%d")}',0)
        for line in lines[1:]:
            if line.startswith('## '): doc.add_heading(line[3:],1)
            elif line.startswith('- '): doc.add_paragraph(line[2:])
        buf=io.BytesIO(); doc.save(buf); buf.seek(0)
        import base64
        return {"markdown":report,"docx_base64":base64.b64encode(buf.read()).decode()}
    except: return {"markdown":report}

# ═══════════ 预测 API ═══════════

@app.get("/api/predict/status")
def api_predict_status():
    if not PREDICT_AVAILABLE:
        return {"trained": False, "model_exists": False, "error": "预测模块未安装 (需要 xgboost, scikit-learn, joblib)"}
    return predict_train.get_status()

@app.get("/api/diagnostics")
def api_diagnostics():
    """系统自检：检查所有组件状态"""
    diag = {
        "server": "ok",
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }
    # K线缓存
    try:
        cache_files = len([f for f in os.listdir(KLINE_CACHE_DIR) if f.endswith('.json')]) if os.path.isdir(KLINE_CACHE_DIR) else 0
        diag["kline_cache"] = {"files": cache_files, "dir": KLINE_CACHE_DIR}
    except Exception:
        diag["kline_cache"] = {"error": "无法访问缓存目录"}

    # ML模型
    if PREDICT_AVAILABLE:
        status = predict_train.get_status()
        diag["ml_model"] = {
            "trained": status.get("model_exists", False),
            "macro_f1": status.get("macro_f1"),
            "outperform_auc": status.get("outperform_auc"),
            "tradable": status.get("model_health", {}).get("tradable", False),
            "reason": status.get("model_health", {}).get("reason", ""),
            "features": status.get("feature_count", 0),
            "samples": status.get("n_samples", 0),
        }
    else:
        diag["ml_model"] = {"error": "预测模块未安装"}

    # Walk-forward
    if WF_AVAILABLE:
        wf_status = predict_wf.get_wf_status(_walkforward_params({}))
        if not wf_status.get("available"):
            wf_status = predict_wf.get_wf_status(None)
        diag["walkforward"] = {
            "available": wf_status.get("available", False),
            "annual_alpha": wf_status.get("summary", {}).get("annual_alpha_pct"),
            "sharpe": wf_status.get("summary", {}).get("sharpe_of_alpha"),
            "win_rate": wf_status.get("summary", {}).get("win_rate"),
        }
    else:
        diag["walkforward"] = {"error": "WF模块未安装"}

    # 雷达
    diag["radar"] = {"available": RADAR_AVAILABLE}

    # 股票池
    universe = load_all_stock_pool()
    diag["universe"] = {"total": universe.get("total", 0), "source": universe.get("source", "unknown")}

    # 判断整体健康
    issues = []
    if diag.get("kline_cache", {}).get("files", 0) < 5:
        issues.append("K线缓存不足，请先浏览行情页面触发数据下载")
    ml = diag.get("ml_model", {})
    if not ml.get("trained"):
        issues.append('ML模型未训练，请点击「训练模型」')
    elif not ml.get("tradable"):
        issues.append(f'模型未达交易门槛: {ml.get("reason", "")}')
    wf = diag.get("walkforward", {})
    if not wf.get("available"):
        issues.append('Walk-Forward验证未完成，请点击「策略验证」')
    diag["issues"] = issues
    diag["healthy"] = len(issues) == 0

    return diag


@app.post("/api/predict/train")
def api_predict_train(data: dict):
    if not PREDICT_AVAILABLE:
        return JSONResponse({"error": "预测模块未安装，请运行: pip install xgboost scikit-learn joblib"}, 500)
    codes = data.get("codes")
    if not codes:
        scope = data.get("universe_scope", "all")
        max_universe = int(data.get("max_universe", 200) or 200)
        universe = resolve_stock_universe(scope, limit=max_universe if max_universe else 200)
        codes = [s["code"] for s in universe.get("items", [])][:200]
    if not codes:
        codes = [s["code"] for s in STOCK_POOL[:50]]
    days = data.get("days", 1200)
    benchmark = data.get("benchmark", "000300")  # 沪深300
    lookahead_days = int(data.get("lookahead_days", 5))
    target_mode = data.get("target_mode", "excess_4class")
    try:
        result = predict_train.train_all_stocks(
            fetch_kline,
            codes,
            days=days,
            benchmark_code=benchmark,
            lookahead_days=lookahead_days,
            target_mode=target_mode,
        )
    except Exception as e:
        return JSONResponse({"error": f"训练失败: {str(e)}"}, 500)
    if "error" in result:
        return JSONResponse(result, 400)
    return result


@app.get("/api/predict/signals")
def api_predict_signals(
    refresh: bool = Query(False),
    scope: str = Query("core"),
    limit: int = Query(0),
    offset: int = Query(0),
):
    if not PREDICT_AVAILABLE:
        return JSONResponse({"error": "预测模块未安装"}, 500)
    universe = resolve_stock_universe(scope, limit=limit, offset=offset)
    items = universe.get("items", [])
    name_map = {s["code"]: s["name"] for s in items}
    codes = [s["code"] for s in items]
    cache_key = _prediction_signals_cache_key(codes)
    now = time.time()
    with PREDICTION_SIGNALS_LOCK:
        cached = PREDICTION_SIGNALS_CACHE.get("data")
        if (cached and PREDICTION_SIGNALS_CACHE.get("key") == cache_key
                and now - PREDICTION_SIGNALS_CACHE.get("ts", 0) <= PREDICTION_SIGNALS_TTL):
            data = apply_decision_overlay_to_payload(dict(cached))
            data["cache"] = {"hit": True, "ttl_seconds": PREDICTION_SIGNALS_TTL, "source": "memory"}
            data["served_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            data["prediction_job"] = _prediction_job_snapshot()
            return data

    snapshot = _load_prediction_snapshot(cache_key)
    if snapshot:
        data = apply_decision_overlay_to_payload(dict(snapshot))
        data["cache"] = {"hit": True, "ttl_seconds": PREDICTION_SIGNALS_TTL, "source": "disk_snapshot"}
        data["served_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        refreshing = False
        if refresh:
            with PREDICTION_JOB_LOCK:
                running = PREDICTION_JOB.get("running")
            if not running:
                threading.Thread(target=_run_prediction_job, args=(cache_key, codes, name_map), daemon=True).start()
            refreshing = True
        data["refreshing"] = refreshing
        data["prediction_job"] = _prediction_job_snapshot()
        return data

    with PREDICTION_JOB_LOCK:
        running = PREDICTION_JOB.get("running")
    if running:
        return JSONResponse({
            "predictions": [],
            "metrics": predict_train.get_status(),
            "updated": "",
            "total": 0,
            "refreshing": True,
            "message": "ML预测正在后台生成，稍后自动刷新",
            "prediction_job": _prediction_job_snapshot(),
            "cache": {"hit": False, "source": "background_running"},
        }, 202)

    if refresh:
        with PREDICTION_JOB_LOCK:
            running = PREDICTION_JOB.get("running")
        if not running:
            threading.Thread(target=_run_prediction_job, args=(cache_key, codes, name_map), daemon=True).start()
        return JSONResponse({
            "predictions": [],
            "metrics": predict_train.get_status(),
            "updated": "",
            "total": 0,
            "refreshing": True,
            "message": "ML预测正在后台生成，稍后自动刷新",
            "prediction_job": _prediction_job_snapshot(),
            "cache": {"hit": False, "source": "background_pending"},
        }, 202)

    data = _build_prediction_signals(codes, name_map, cache_key)
    if data.get("error"):
        return JSONResponse(data, 400)
    with PREDICTION_SIGNALS_LOCK:
        PREDICTION_SIGNALS_CACHE.update({"key": cache_key, "ts": time.time(), "data": data})
    _write_json_file(PREDICTION_SIGNALS_PATH, data)
    return data


@app.get("/api/predict/signal")
def api_predict_signal(code: str = Query(...)):
    if not PREDICT_AVAILABLE:
        return JSONResponse({"error": "预测模块未安装"}, 500)
    model = predict_train.load_model()
    if model is None:
        return JSONResponse({"error": "模型未训练，请先 POST /api/predict/train"}, 400)
    name_map = stock_name_map("all")
    result = predict_train.predict_single(
        fetch_kline,
        code,
        benchmark_code="000300",
        model=model,
        feature_context=predict_train._load_feature_context([code]),
        horizon_days=predict_train._model_lookahead_days(),
    )
    if "error" in result:
        return JSONResponse(result, 400)
    result["name"] = name_map.get(code, code)
    result = apply_decision_overlay(result)
    return result


@app.get("/api/predict/forecast/{code}")
def api_predict_forecast(code: str, horizon: int = Query(5)):
    if not PREDICT_AVAILABLE:
        return JSONResponse({"error": "预测模块未安装"}, 500)
    model = predict_train.load_model()
    if model is None:
        return JSONResponse({"error": "模型未训练，请先 POST /api/predict/train"}, 400)
    kl = fetch_kline_prediction(code, FORECAST_HISTORY_DAYS) or fetch_kline(code, FORECAST_HISTORY_DAYS)
    if not kl or len(kl) < 80:
        return JSONResponse({"error": f"{code} 数据不足，无法生成预测K线"}, 400)
    result = predict_train.predict_single(
        fetch_kline_prediction,
        code,
        benchmark_code="000300",
        model=model,
        feature_context=predict_train._load_feature_context([code]),
        horizon_days=predict_train._model_lookahead_days(),
        stock_kl=kl,
        benchmark_kl=fetch_kline_prediction("000300", FORECAST_HISTORY_DAYS),
    )
    if "error" in result:
        return JSONResponse(result, 400)
    result["name"] = stock_name(code)
    result = apply_decision_overlay(result)
    forecast = build_forecast_path(kl, result, horizon=horizon)
    return {
        "code": code,
        "name": result["name"],
        "prediction": result,
        "forecast": forecast,
    }

# ═══════════ Walk-Forward 回测 ═══════════

def _walkforward_params(data=None):
    data = data or {}
    return {
        "codes": data.get("codes") or [s["code"] for s in STOCK_POOL[:50]],
        "benchmark_code": data.get("benchmark_code", "000300"),
        "train_years": int(data.get("train_years", 5)),
        "step_months": int(data.get("step_months", 3)),
        "lookahead_days": int(data.get("lookahead_days", 5)),
        "top_n": int(data.get("top_n", 8)),
        "max_bars": int(data.get("max_bars", 2500)),
        "commission": float(data.get("commission", 0.0003)),
        "slippage": float(data.get("slippage", 0.0005)),
        "stamp_tax": float(data.get("stamp_tax", 0.0005)),
        "spread": float(data.get("spread", 0.0002)),
        "max_train_samples": int(data.get("max_train_samples", 0)),
        "n_estimators": int(data.get("n_estimators", 200)),
    }


@app.on_event("startup")
def start_walkforward_cache_warmup():
    if not WF_AVAILABLE:
        return
    # 已有结果就不重跑（检查主文件，不依赖特定参数）
    if os.path.exists(predict_wf.WF_RESULTS_PATH):
        return

    def warm_default_walkforward():
        with WF_JOB_LOCK:
            if WF_JOB.get("running"):
                return
        _run_walkforward_job(_walkforward_params({}))

    threading.Thread(target=warm_default_walkforward, daemon=True).start()


@app.post("/api/predict/walkforward")
def api_walkforward(data: dict):
    if not WF_AVAILABLE:
        return JSONResponse({"error": "Walk-forward 模块未安装"}, 500)
    with WF_JOB_LOCK:
        if WF_JOB.get("running"):
            return JSONResponse({
                "started": False,
                "running": True,
                "message": "策略验证已经在后台运行",
                "job": dict(WF_JOB),
            }, 202)

    params = _walkforward_params(data)
    if not data.get("force"):
        cached = predict_wf.get_wf_status(params)
        if cached.get("available"):
            cached["cached"] = True
            cached["started"] = False
            cached["running"] = False
            cached["message"] = "已返回最近一次Walk-Forward结果；传 force=true 才会后台重跑完整验证"
            cached["job"] = _wf_job_snapshot()
            return cached

    if data.get("sync") is True:
        result = predict_wf.run_walkforward(fetch_kline, **params)
        if "error" in result:
            return JSONResponse(result, 400)
        return result

    thread = threading.Thread(target=_run_walkforward_job, args=(params,), daemon=True)
    thread.start()
    return JSONResponse({
        "started": True,
        "running": True,
        "message": "策略验证已在后台启动，轮询 /api/predict/walkforward/status 查看进度",
        "params": params,
    }, 202)


@app.get("/api/predict/walkforward/status")
def api_walkforward_status(default: bool = Query(False)):
    if not WF_AVAILABLE:
        return JSONResponse({"available": False}, 200)
    # 先尝试精确匹配, 不行就返回最新结果
    if default:
        status = predict_wf.get_wf_status(_walkforward_params({}))
        if not status.get("available"):
            status = predict_wf.get_wf_status(None)  # fallback to main results file
    else:
        status = predict_wf.get_wf_status(None)
    status["job"] = _wf_job_snapshot()
    return status


# ═══════════ 雷达 API ═══════════

@app.get("/api/radar/news/{code}")
def api_radar_news(code: str):
    if not RADAR_AVAILABLE:
        return JSONResponse({"error": "雷达模块未安装"}, 500)
    name = next((s["name"] for s in STOCK_POOL if s["code"] == code), code)
    news_data = radar_news.fetch_news_for_code(code, stock_name=name, limit=8)
    sentiment = radar_sentiment.analyze_sentiment(news_data.get("news", []), code, name)
    sentiment.setdefault("news_quality", news_data.get("news_quality", "unknown"))
    news_data["sentiment"] = sentiment
    news_data.setdefault("diagnostics", {})["sentiment_fallback_used"] = bool(sentiment.get("_fallback"))
    return news_data


@app.get("/api/radar/sentiment/{code}")
def api_radar_sentiment(code: str):
    if not RADAR_AVAILABLE:
        return JSONResponse({"error": "雷达模块未安装"}, 500)
    name = next((s["name"] for s in STOCK_POOL if s["code"] == code), code)
    news_data = radar_news.fetch_news_for_code(code, stock_name=name, limit=6)
    sentiment = radar_sentiment.analyze_sentiment(news_data.get("news", []), code, name)
    sentiment.setdefault("news_quality", news_data.get("news_quality", "unknown"))
    return sentiment


@app.post("/api/radar/scan")
def api_radar_scan(data: dict):
    if not RADAR_AVAILABLE:
        return JSONResponse({"error": "雷达模块未安装"}, 500)
    codes = (data.get("codes") or [s["code"] for s in STOCK_POOL[:20]])
    name_map = {s["code"]: s["name"] for s in STOCK_POOL}
    results = []
    for code in codes:
        name = name_map.get(code, code)
        news_data = radar_news.fetch_news_for_code(code, stock_name=name, limit=5)
        sentiment = radar_sentiment.analyze_sentiment(news_data.get("news", []), code, name)
        if abs(sentiment.get("score", 0)) > 0.2:
            results.append({
                "code": code, "name": name,
                "sentiment": sentiment,
                "news_count": news_data.get("total", 0),
            })
    results.sort(key=lambda r: abs(r["sentiment"]["score"]), reverse=True)
    return {"signals": results, "total": len(results), "scanned": len(codes),
            "scanned_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")}


@app.get("/api/radar/risk/{code}")
def api_radar_risk(code: str):
    if not RADAR_AVAILABLE:
        return JSONResponse({"error": "雷达模块未安装"}, 500)
    name = next((s["name"] for s in STOCK_POOL if s["code"] == code), code)
    news_data = radar_news.fetch_news_for_code(code, stock_name=name, limit=5)
    sentiment = radar_sentiment.analyze_sentiment(news_data.get("news", []), code, name)
    sentiment.setdefault("news_quality", news_data.get("news_quality", "unknown"))

    # Get ML prediction
    kl = fetch_kline_prediction(code, 500) or fetch_kline(code, 120)
    if PREDICT_AVAILABLE:
        model = predict_train.load_model()
        if model:
            benchmark_kl = fetch_kline_prediction("000300", 500) or fetch_kline("000300", 500)
            ml_pred = predict_train.predict_single(
                fetch_kline_prediction,
                code,
                benchmark_code="000300",
                model=model,
                feature_context=predict_train._load_feature_context([code]),
                stock_kl=kl,
                benchmark_kl=benchmark_kl,
                horizon_days=predict_train._model_lookahead_days(),
            )
        else:
            ml_pred = {"prob_up": 0.5, "signal": "中性", "confidence": "低"}
    else:
        ml_pred = {"prob_up": 0.5, "signal": "中性", "confidence": "低"}

    k_feats = risk_kline_features(code, kl, ml_pred)
    risk = radar_risk.assess_risk(code, sentiment, ml_pred, k_feats)
    risk["name"] = name
    risk["sentiment"] = sentiment
    risk["news_diagnostics"] = news_data.get("diagnostics", {})
    return risk


@app.get("/api/radar/risks")
def api_radar_risks():
    if not RADAR_AVAILABLE:
        return JSONResponse({"error": "雷达模块未安装"}, 500)
    name_map = {s["code"]: s["name"] for s in STOCK_POOL}
    codes = [s["code"] for s in STOCK_POOL[:30]]
    model = predict_train.load_model() if PREDICT_AVAILABLE else None
    feature_context = predict_train._load_feature_context(codes) if model else None
    horizon_days = predict_train._model_lookahead_days() if model else 5
    benchmark_kl = (fetch_kline_prediction("000300", 500) or fetch_kline("000300", 500)) if model else None
    def analyze_one(code):
        try:
            name = name_map.get(code, code)
            news_data = radar_news.fetch_news_for_code(code, stock_name=name, limit=3)
            sentiment = radar_sentiment.analyze_sentiment(news_data.get("news", []), code, name)
            sentiment.setdefault("news_quality", news_data.get("news_quality", "unknown"))
            kl = fetch_kline_prediction(code, 500) or fetch_kline(code, 120)
            ml_pred = predict_train.predict_single(
                fetch_kline_prediction,
                code,
                benchmark_code="000300",
                model=model,
                feature_context=feature_context,
                horizon_days=horizon_days,
                stock_kl=kl,
                benchmark_kl=benchmark_kl,
            ) if model else {"prob_up": 0.5, "signal": "中性", "confidence": "低"}
            k_feats = risk_kline_features(code, kl, ml_pred)
            risk = radar_risk.assess_risk(code, sentiment, ml_pred, k_feats)
            return {
                "code": code, "name": name,
                "sentiment_score": sentiment.get("score", 0),
                "sentiment_dir": sentiment.get("direction", "neutral"),
                "news_quality": sentiment.get("news_quality", "unknown"),
                "news_count": sentiment.get("news_count", news_data.get("total", 0)),
                "outperform_prob": ml_pred.get("outperform_prob", ml_pred.get("prob_up", 0.5)),
                "risk_level": risk["level"],
                "risk_type": risk["type"],
                "risk_score": risk["score"],
                "top_reason": risk["reasons"][0] if risk.get("reasons") else "",
            }
        except Exception:
            return None

    results = []
    with ThreadPoolExecutor(max_workers=min(8, max(1, len(codes)))) as pool:
        for future in as_completed([pool.submit(analyze_one, code) for code in codes]):
            item = future.result()
            if item:
                results.append(item)
    level_rank = {"极高": 4, "高": 3, "中": 2, "低": 1}
    results.sort(key=lambda r: (level_rank.get(r.get("risk_level"), 0), r.get("risk_score", 0)), reverse=True)

    bullish = sum(1 for r in results if (r.get("sentiment_score") or 0) > 0.2)
    bearish = sum(1 for r in results if (r.get("sentiment_score") or 0) < -0.2)
    high_risk = sum(1 for r in results if r.get("risk_level") in ("高", "极高"))

    low_quality = sum(1 for r in results if r.get("news_quality") in ("none", "low"))

    return {
        "stocks": results,
        "summary": {"bullish": bullish, "bearish": bearish, "neutral": len(results) - bullish - bearish,
                    "high_risk": high_risk, "low_news_quality": low_quality, "total": len(results)},
        "scanned_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }


@app.get("/api/predict/adjusted/{code}")
def api_predict_adjusted_code(code: str):
    if not PREDICT_AVAILABLE:
        return JSONResponse({"error": "预测模块未安装"}, 500)
    name = next((s["name"] for s in STOCK_POOL if s["code"] == code), code)
    model = predict_train.load_model()
    feature_context = predict_train._load_feature_context([code]) if model else None
    kl = fetch_kline_prediction(code, 500) or fetch_kline(code, 120)
    benchmark_kl = fetch_kline_prediction("000300", 500) or fetch_kline("000300", 500) if model else None
    ml_pred = predict_train.predict_single(
        fetch_kline_prediction,
        code,
        benchmark_code="000300",
        model=model,
        feature_context=feature_context,
        stock_kl=kl,
        benchmark_kl=benchmark_kl,
        horizon_days=predict_train._model_lookahead_days(),
    ) if model else {"prob_up": 0.5, "signal": "中性", "confidence": "低"}

    sentiment = {"score": 0, "confidence": 0}
    if RADAR_AVAILABLE:
        try:
            news_data = radar_news.fetch_news_for_code(code, stock_name=name, limit=3)
            sentiment = radar_sentiment.analyze_sentiment(news_data.get("news", []), code, name)
            sentiment.setdefault("news_quality", news_data.get("news_quality", "unknown"))
        except Exception:
            pass

    risk_level = "低"
    if RADAR_AVAILABLE:
        try:
            k_feats = risk_kline_features(code, kl, ml_pred)
            risk = radar_risk.assess_risk(code, sentiment, ml_pred, k_feats)
            risk_level = risk["level"]
        except Exception:
            pass

    adjusted = radar_scorer.adjusted_probability(
        ml_pred.get("prob_up", 0.5),
        sentiment.get("score", 0),
        sentiment.get("confidence", 0),
        risk_level,
    )
    return {
        "code": code, "name": name,
        "ml_prediction": ml_pred,
        "sentiment": sentiment,
        "risk_level": risk_level,
        "adjusted": adjusted,
    }


@app.get("/api/predict/adjusted")
def api_predict_adjusted_all():
    name_map = {s["code"]: s["name"] for s in STOCK_POOL}
    codes = [s["code"] for s in STOCK_POOL[:30]]
    model = predict_train.load_model() if PREDICT_AVAILABLE else None
    feature_context = predict_train._load_feature_context(codes) if model else None
    horizon_days = predict_train._model_lookahead_days() if model else 5
    benchmark_kl = (fetch_kline_prediction("000300", 500) or fetch_kline("000300", 500)) if model else None
    def analyze_one(code):
        try:
            kl = fetch_kline_prediction(code, 500) or fetch_kline(code, 120)
            ml_pred = predict_train.predict_single(
                fetch_kline_prediction,
                code,
                benchmark_code="000300",
                model=model,
                feature_context=feature_context,
                horizon_days=horizon_days,
                stock_kl=kl,
                benchmark_kl=benchmark_kl,
            ) if model else {"prob_up": 0.5, "signal": "中性", "confidence": "低"}
            sentiment = {"score": 0, "confidence": 0}
            if RADAR_AVAILABLE:
                try:
                    news_data = radar_news.fetch_news_for_code(code, stock_name=name_map.get(code, code), limit=3)
                    sentiment = radar_sentiment.analyze_sentiment(news_data.get("news", []), code, name_map.get(code, code))
                    sentiment.setdefault("news_quality", news_data.get("news_quality", "unknown"))
                except Exception:
                    pass
            risk_level = "低"
            if RADAR_AVAILABLE:
                try:
                    k_feats = risk_kline_features(code, kl, ml_pred)
                    risk = radar_risk.assess_risk(code, sentiment, ml_pred, k_feats)
                    risk_level = risk["level"]
                except Exception:
                    pass
            adjusted = radar_scorer.adjusted_probability(
                ml_pred.get("prob_up", 0.5), sentiment.get("score", 0), sentiment.get("confidence", 0), risk_level)
            return {"code": code, "name": name_map.get(code, code),
                    "prob_xgb": ml_pred.get("prob_up", 0.5),
                    "outperform_prob": ml_pred.get("outperform_prob", ml_pred.get("prob_up", 0.5)),
                    "sentiment_score": sentiment.get("score", 0),
                    "risk_level": risk_level,
                    "prob_adjusted": adjusted["prob_adjusted"],
                    "signal": adjusted["signal"],
                    "sentiment_offset": adjusted["sentiment_offset"]}
        except Exception:
            return None

    results = []
    with ThreadPoolExecutor(max_workers=min(8, max(1, len(codes)))) as pool:
        for future in as_completed([pool.submit(analyze_one, code) for code in codes]):
            item = future.result()
            if item:
                results.append(item)
    results.sort(key=lambda r: abs(r["prob_adjusted"] - 0.5), reverse=True)
    return {"predictions": results, "total": len(results),
            "updated": datetime.now().strftime("%Y-%m-%d %H:%M:%S")}


app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/")
def index(): return FileResponse("static/index.html")

if __name__=="__main__":
    import uvicorn
    print("🟢 QuantDesk 启动 → http://localhost:8888")
    uvicorn.run(app, host="0.0.0.0", port=8888, log_level="warning")
